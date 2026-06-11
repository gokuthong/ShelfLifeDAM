# -*- coding: ascii -*-
"""
Builds BryanTeyKaiYuan_5011CEM_CW1_Report.docx (Individual Progressive Report).
All facts come from 06_Working_Files/REPORT_FACTS_PACK.md.
References restricted to 06_Working_Files/verified_references.md (Bryan allocation + shared core).
Regenerate with:  python -X utf8 build_bryan_report.py
"""

import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\ASUS\Downloads\07_Team_Handoff\07_Team_Handoff"
OUT_DIR = os.path.join(BASE, "08_Final_Submission")
OUT_DOCX = os.path.join(OUT_DIR, "BryanTeyKaiYuan_5011CEM_CW1_Report.docx")

IMG = {
    "eda_target": os.path.join(BASE, "04_Diagrams", "EDA_Plots", "eda_target.png"),
    "eda_ext_vs_target": os.path.join(BASE, "04_Diagrams", "EDA_Plots", "eda_ext_vs_target.png"),
    "dfd_level0": os.path.join(BASE, "04_Diagrams", "System_Diagrams", "dfd_level0.png"),
    "dfd_level1": os.path.join(BASE, "04_Diagrams", "System_Diagrams", "dfd_level1.png"),
    "erd": os.path.join(BASE, "04_Diagrams", "System_Diagrams", "erd.png"),
    "gantt": os.path.join(BASE, "06_Working_Files", "gantt_homecredit.png"),
}

FONT = "Times New Roman"


# ----------------------------------------------------------------------------
# style helpers
# ----------------------------------------------------------------------------

def set_style_font(style, name=FONT, size=12, bold=None, color=RGBColor(0, 0, 0)):
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=12, color=None)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for sname, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[sname]
        set_style_font(st, size=size, bold=True)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def para(doc, text, align=None, bold=False, size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def page_break(doc):
    doc.add_page_break()


def add_figure(doc, key, caption):
    path = IMG[key]
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(path, width=Inches(6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(11)


def add_caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def add_table(doc, header, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(header):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(htxt)
        r.bold = True
        r.font.size = Pt(11)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(val))
            r.font.size = Pt(11)
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(len(rows) + 1):
                t.cell(i, j).width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    placeholder = p.add_run("Table of contents entries appear here. Select this text and press F9 to update the field.")
    placeholder.font.size = Pt(11)
    run_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


# ----------------------------------------------------------------------------
# cover sheet
# ----------------------------------------------------------------------------

def build_cover(doc):
    para(doc, "INTI International College Penang", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    para(doc, "School of Computing", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    para(doc, "3+0 Bachelor of Science (Hons) in Computer Science, in collaboration with Coventry University, UK",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "Coursework cover sheet", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    para(doc, "Section A - To be completed by the student", bold=True)

    rows = [
        ("Full Name", "Bryan Tey Kai Yuan"),
        ("CU Student ID Number", "P23015693"),
        ("Semester", "April 2026"),
        ("Lecturer", "Ms. Vimala Doraisamy"),
        ("Module Code and Title", "5011CEM Big Data Programming Project"),
        ("Assignment No. / Title", "Home Credit Loan-Default Prediction - 100% of Module Mark"),
        ("Hand out date", "Week 1"),
        ("Due date", "Individual Progress Report - WK10 / Final Report - WK12 / VIVA - WK13 & 14"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        c0 = t.cell(i, 0)
        c0.text = ""
        p0 = c0.paragraphs[0]
        p0.paragraph_format.line_spacing = 1.0
        r0 = p0.add_run(k)
        r0.bold = True
        c1 = t.cell(i, 1)
        c1.text = ""
        p1 = c1.paragraphs[0]
        p1.paragraph_format.line_spacing = 1.0
        p1.add_run(v)
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)

    doc.add_paragraph()
    para(doc, "Penalties:", bold=True, space_after=0)
    para(doc, "No late work will be accepted. If you are unable to submit coursework on time due to "
              "extenuating circumstances, you may be eligible for an extension. Please consult the lecturer.")
    para(doc, "Declaration:", bold=True, space_after=0)
    para(doc, "I/we the undersigned confirm that I/we have read and agree to abide by the University "
              "regulations on plagiarism and cheating and Faculty coursework policies and procedures. "
              "I/we confirm that this piece of work is my/our own. I/we consent to appropriate storage "
              "of our work for plagiarism checking.")
    para(doc, "Signature(s): Bryan Tey Kai Yuan")
    doc.add_paragraph()
    para(doc, "Individual Progressive Report", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    para(doc, "Home Credit Loan-Default Prediction", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)


# ----------------------------------------------------------------------------
# body content
# ----------------------------------------------------------------------------

def build_chapter1(doc):
    h1(doc, "Chapter 1: Introduction")

    h2(doc, "1.1 Problem Background")
    para(doc,
         "Lending is the core business of banks and consumer finance companies, but every approved loan "
         "carries the possibility that the borrower fails to repay. Home Credit is a consumer finance "
         "provider that focuses on applicants with little or no formal credit history, so the usual credit "
         "bureau checks are often thin or absent for its customers. When a lender cannot judge an applicant "
         "accurately, two expensive mistakes happen. Firstly, an applicant who later defaults may be "
         "approved, and the unpaid balance becomes a direct write-off against the loan book. Secondly, a "
         "reliable applicant may be rejected, which loses interest income and pushes a genuine customer "
         "towards a competitor. My team and I selected loan default prediction in the banking and financial "
         "sector as our project because it is a realistic big data problem where analytics can directly "
         "improve a business decision that is made thousands of times every day. The decision is also "
         "asymmetric, because a single written off loan can wipe out the interest margin earned from many "
         "good loans, which is why lenders care more about catching likely defaulters than about a tidy "
         "overall accuracy number.")
    para(doc,
         "The project uses the Home Credit Default Risk dataset that was released for a public Kaggle "
         "competition in 2018. The core application table contains 307,511 historical loan applications "
         "described by 122 attributes, and only 8.07 percent of these applications (24,825 cases) ended in "
         "payment difficulty. This scale and imbalance reflect a genuine retail credit portfolio rather "
         "than a toy classroom sample. Research on big data in finance reports that data driven risk "
         "assessment improves decision quality and is changing how institutions manage credit risk (Hasan, "
         "Popp and Olah, 2020). Furthermore, banks that apply analytics across fraud detection, customer "
         "behaviour and risk management gain measurable advantages over those that do not (Srivastava and "
         "Gopalkrishnan, 2015). In a nutshell, a model that ranks applicants by default risk, supported by "
         "clear visual reporting, gives credit analysts a practical tool for protecting portfolio quality.")

    h2(doc, "1.2 Scope and Team Responsibilities")
    para(doc,
         "This individual progressive report covers the design and planning stage of the project together "
         "with early exploratory evidence from the data. The implementation scope for the current phase is "
         "the single application_train table. The six relational child tables supplied with the competition "
         "are documented as a future extension and are not part of the planned first model. The project is "
         "carried out by a two person team. I am responsible for the exploratory data analysis, the "
         "Logistic Regression and Random Forest models, the dashboard Overview and Live Prediction pages, "
         "and the assembly of the group final report. My teammate Thong Wai Kit leads the data loading and "
         "preprocessing modules, the XGBoost model with its class imbalance experiment, the evaluation "
         "module, the dashboard Model Comparison page, and the architecture documentation. Joint decisions, "
         "such as the choice of dataset and the shared evaluation metrics, were made by my team and I "
         "together at the start of the semester.")

    h2(doc, "1.3 Goals and Objectives")
    para(doc,
         "The overall aim is to design a reproducible analytics pipeline that predicts the probability of "
         "loan default and presents the result in a form that a credit analyst can act on. Four objectives "
         "support this aim. Firstly, I want to explore the application data and identify which applicant "
         "attributes are most strongly associated with default. Secondly, I want to confirm these "
         "associations with formal statistical tests instead of relying on visual impressions alone. "
         "Thirdly, the project will design and later train at least two classification models, a Logistic "
         "Regression baseline and a Random Forest ensemble, which will then be compared against the XGBoost "
         "model prepared by my teammate. Lastly, the project will deliver a Streamlit dashboard that "
         "communicates risk scores as simple risk bands so that non technical staff can interpret the "
         "output without statistical training.")

    h2(doc, "1.4 Success Criteria")
    para(doc,
         "Success for this progressive stage is defined in measurable terms. The statistical analysis must "
         "confirm or reject each hypothesis at the five percent significance level. The preprocessing "
         "design must avoid data leakage by fitting every transformation on training data only. For the "
         "modelling stage that follows, the models will be judged mainly on recall for the defaulter class "
         "and on ranking quality measured by ROC AUC, because plain accuracy is misleading when only about "
         "eight percent of applications default. A model that simply approves everybody would already score "
         "about 92 percent accuracy while catching no defaulters at all. In addition, the project must "
         "remain fully reproducible, with version controlled code, automated unit tests and a fixed random "
         "seed for the train and test split. Lastly, the progressive deliverables themselves act as "
         "criteria: the statistical evidence, the architecture diagrams and the plans written in this "
         "report must be complete enough that the remaining weeks can concentrate on model training and "
         "the dashboard build without revisiting design decisions.")

    h2(doc, "1.5 Report Structure")
    para(doc,
         "The remainder of this report is organised as follows. Chapter 2 reviews the literature on credit "
         "scoring, compares candidate techniques and derives the research questions and hypotheses. Chapter "
         "3 describes the data source and schema, the data quality issues found during early exploration, "
         "the preprocessing plan, the system architecture with its contextual diagram, and the data storage "
         "options. Chapter 4 presents the analysis tools, the hypothesis test results obtained so far, the "
         "justification of the chosen algorithms, and the implementation and evaluation plan. The reference "
         "list follows Chapter 4. The appendices contain the project Gantt chart, the level 1 data flow "
         "diagram and the entity relationship diagram, and the final page is reserved for the Turnitin "
         "similarity receipt.")


def build_chapter2(doc):
    h1(doc, "Chapter 2: Literature Review")

    h2(doc, "2.1 Credit Scoring Applications in Practice")
    para(doc,
         "Credit scoring is one of the oldest industrial applications of statistical classification. Hand "
         "and Henley (1997) reviewed the methods used by lenders and explained why logistic regression "
         "became the industry baseline. Its coefficients translate directly into scorecard points, so an "
         "analyst can see exactly how each attribute moves an applicant up or down, and a regulator can "
         "audit the decision rule. That property still matters today, because a lender must be able to "
         "explain why an application was declined.")
    para(doc,
         "Later benchmark studies asked whether more complex algorithms are worth the extra opacity. "
         "Baesens et al. (2003) compared seventeen classifiers on eight real credit datasets and found that "
         "the performance differences were often small, with simple linear models remaining competitive. "
         "On the other hand, the updated benchmark by Lessmann et al. (2015), which covered forty one "
         "classifiers, showed that ensemble methods such as Random Forest consistently beat logistic "
         "regression on ranking quality. The improvement per application is modest, but across a portfolio "
         "of hundreds of thousands of loans even a small gain in ranking accuracy changes the amount of "
         "money a lender saves. These two studies together justify the comparison plan in this project, "
         "which keeps an interpretable baseline and tests whether ensembles add real value on our data. "
         "For a lender, the practical question is not which algorithm tops a leaderboard but whether the "
         "extra discrimination converts into fewer write-offs without turning away more good customers, "
         "and that framing guides how the models in this project will be compared.")
    para(doc,
         "Most relevant of all, Yang et al. (2025) modelled the same Home Credit dataset used in this "
         "project. They trained logistic regression, Random Forest and boosting models, found that the "
         "ensembles won, and showed through SHAP analysis that the three external credit scores, the "
         "EXT_SOURCE attributes, dominate every other feature. This gives my team and I a clear prior "
         "expectation for our own exploratory work, and it is the main reason the first hypothesis in "
         "Section 2.4 concentrates on the external scores.")

    h2(doc, "2.2 Techniques and Algorithms Compared")
    para(doc,
         "Logistic regression models the log odds of default as a weighted sum of the input attributes. It "
         "trains quickly, behaves predictably on wide one-hot encoded data, and produces coefficients that "
         "support scorecard style explanation (Hand and Henley, 1997). Its weakness is that it can only "
         "draw linear decision boundaries, so interactions such as young age combined with short employment "
         "must be engineered manually.")
    para(doc,
         "Random Forest addresses that weakness by averaging many decorrelated decision trees, each grown "
         "on a bootstrap sample with random feature selection at every split (Breiman, 2001). The ensemble "
         "captures non linear effects and interactions automatically and is robust to outliers and noisy "
         "attributes. The cost is that a forest of hundreds of trees cannot be read the way a scorecard "
         "can. SHAP values offer a repair for this, because they attribute each individual prediction to "
         "the input features in a consistent game theoretic way (Lundberg and Lee, 2017). The project "
         "therefore plans to pair the ensemble models with importance and SHAP style explanation rather "
         "than treating them as black boxes.")
    para(doc,
         "Class imbalance is the third recurring theme in the literature. With roughly 92 repaid loans for "
         "every 8 defaults, a naive learner ignores the minority class. One classic remedy is SMOTE, which "
         "generates synthetic minority examples by interpolating between neighbouring defaulters (Chawla et "
         "al., 2002). An alternative is cost sensitive learning, where class weights make each defaulter "
         "count more during training. My team and I plan to apply exactly one correction per model, using "
         "balanced class weights for Logistic Regression and Random Forest, because stacking several "
         "corrections tends to inflate recall while destroying precision. My teammate will test this "
         "design choice empirically in his XGBoost imbalance experiment.")

    h2(doc, "2.3 Big Data Challenges in Banking")
    para(doc,
         "The banking sector generates data at a volume and variety that strains traditional tools. Hasan, "
         "Popp and Olah (2020) found that big data influences financial analysis and risk management most "
         "strongly, while also noting persistent challenges in data quality, skills and infrastructure. "
         "Srivastava and Gopalkrishnan (2015) reached a similar conclusion for retail banks, describing how "
         "transaction histories, application records and behavioural data feed fraud detection and credit "
         "decisions. Both papers argue that the bottleneck is rarely the algorithm; it is the ability to "
         "store, clean and serve large datasets reliably. This project meets the same issues at a smaller "
         "scale, since the single application table is a 166 MB file with 307,511 rows and the full "
         "relational set is about 2.7 GB across seven tables, which already makes naive reloading painful. "
         "The project responds at its own scale with a parquet cache and one consolidated metrics file, "
         "which are small but genuine versions of the storage and data serving discipline that the "
         "literature recommends.")
    para(doc,
         "Besides the technical challenges, credit models face legal and ethical constraints. Article 22 of "
         "the GDPR gives a person the right not to be subject to a decision based solely on automated "
         "processing where that decision has significant effects, and credit approval is a textbook example "
         "(European Parliament and Council, 2016). Goodman and Flaxman (2017) interpret the regulation as "
         "creating a right to explanation, which pushes lenders towards models whose output can be "
         "justified to the affected customer. Moreover, fairness research shows that credit models can "
         "discriminate through attributes such as gender even when accuracy looks healthy, and that "
         "fairness and profit must be balanced deliberately (Kozodoi, Jacob and Lessmann, 2022). These "
         "findings shape two design decisions in this project: the dashboard keeps a human analyst in the "
         "decision loop, and demographic attributes are treated as material for an ethics discussion rather "
         "than as free predictive fuel.")

    h2(doc, "2.4 Research Questions and Hypotheses")
    para(doc,
         "The review above leads to three research questions framed from the lender point of view. RQ1 "
         "asks which applicant attributes most strongly indicate default risk in the Home Credit portfolio. "
         "RQ2 asks whether interpretable models, namely Logistic Regression and Random Forest, can reach a "
         "level of defaulter recall that is acceptable for lender use. RQ3 asks how model output should be "
         "presented so that credit analysts can use it in daily approval decisions. RQ1 matters because "
         "attribute level evidence guides both feature selection and the explanations analysts give to "
         "customers. RQ2 matters because an opaque model that cannot be defended to a regulator is "
         "unusable in credit, whatever its accuracy. RQ3 matters because a model only changes outcomes "
         "when its output reaches the person who approves the loan in a form that person trusts.")
    para(doc,
         "Three testable hypotheses follow from these questions. H1 states that the external credit scores "
         "EXT_SOURCE_1, EXT_SOURCE_2 and EXT_SOURCE_3 correlate negatively and significantly with default, "
         "in line with the SHAP findings of Yang et al. (2025). H2 states that socioeconomic categorical "
         "attributes, specifically income type, education level, gender, family status and contract type, "
         "are significantly associated with default status. H3 states that ensemble models will achieve "
         "higher ranking quality than Logistic Regression, while Logistic Regression with balanced class "
         "weights will remain competitive on defaulter recall, as the benchmark literature suggests "
         "(Lessmann et al., 2015). H1 and H2 are tested in Chapter 4 with point biserial correlation and "
         "chi square tests on the real data. H3 concerns model performance and will be evaluated in the "
         "final group report once all three models are trained.")


def build_chapter3(doc):
    h1(doc, "Chapter 3: Design and Planning of Project / Data Analysis")

    h2(doc, "3.1 Data Source, Schema and Structure")
    para(doc,
         "The data source is the Home Credit Default Risk dataset published on Kaggle by Home Credit Group "
         "in 2018 for a public competition. The core file, application_train.csv, holds 307,511 rows and "
         "122 columns in 166,133,370 bytes, with one row per loan application. The label column TARGET "
         "equals 1 when the client had payment difficulty and 0 when the loan was repaid, and the overall "
         "default rate is 8.07 percent. The data is real, anonymised consumer loan data with no direct "
         "identifiers, and it comfortably satisfies the assignment requirement of at least 100,000 records. "
         "A separate data dictionary, HomeCredit_columns_description.csv, documents every column.")
    para(doc,
         "The attributes mix numerical and categorical types. Numerical examples include AMT_INCOME_TOTAL, "
         "AMT_CREDIT, AMT_ANNUITY, DAYS_BIRTH, DAYS_EMPLOYED and the three external scores EXT_SOURCE_1 to "
         "EXT_SOURCE_3. Categorical examples include NAME_CONTRACT_TYPE, CODE_GENDER, NAME_EDUCATION_TYPE, "
         "NAME_INCOME_TYPE and NAME_FAMILY_STATUS, alongside a long block of binary document flags. Six "
         "child tables, bureau with bureau_balance, previous_application, POS_CASH_balance, "
         "installments_payments and credit_card_balance, join to the application table through SK_ID_CURR, "
         "and their structure is shown in the entity relationship diagram in Appendix C (Figure C.1). The "
         "current design deliberately limits modelling to the application table, and the child tables are "
         "recorded as the documented extension path for future feature engineering.")
    para(doc,
         "Figure 3.1 shows the class balance explored during early EDA. The 91.9 percent to 8.1 percent "
         "split is the single most important structural fact about this dataset, because it dictates the "
         "choice of evaluation metrics and the class weighting strategy discussed in Chapter 4.")
    add_figure(doc, "eda_target",
               "Figure 3.1: Class balance of the TARGET variable, showing 91.9% repaid against 8.1% default")

    h2(doc, "3.2 Data Quality Issues and Preprocessing Plan")
    para(doc,
         "Early exploration uncovered three quality issues that the preprocessing design must handle. "
         "Firstly, missing values are widespread: roughly 24 percent of all cells are empty, and the sparse "
         "building survey block (the COMMONAREA_*, NONLIVINGAPARTMENTS_* and YEARS_BUILD_* style columns) "
         "is mostly blank. Secondly, the DAYS_EMPLOYED column contains the sentinel value 365243, which "
         "corresponds to roughly one thousand years of employment and actually marks pensioners and "
         "unemployed applicants. This sentinel affects 55,374 rows, about 18.0 percent of the data, and is "
         "replaced with a missing value at load time so it cannot distort any statistic. Thirdly, a "
         "duplicate row check is included in the cleaning step, although the real data contains zero exact "
         "duplicates.")
    para(doc,
         "The treatment of missing data follows the survey by Emmanuel et al. (2021), which distinguishes "
         "data that is missing completely at random from data whose absence carries meaning. The building "
         "survey block falls in the second group, since whole sections were simply never collected for most "
         "applicants, so the plan drops any column with 60 percent or more missing values rather than "
         "fabricating them. The remaining numeric gaps will be filled with the column median, which is "
         "robust to the heavy right skew seen in the income and credit amounts, and categorical gaps will "
         "be filled with the column mode. The FLAG_DOCUMENT_* indicators are dropped as near useless "
         "administrative flags, and the identifier SK_ID_CURR and the label TARGET are excluded from the "
         "feature set.")
    para(doc,
         "The plan also adds six engineered ratios with direct business meaning: AGE_YEARS, YEARS_EMPLOYED, "
         "CREDIT_INCOME_RATIO, ANNUITY_INCOME_RATIO, CREDIT_TERM and EMPLOYED_AGE_RATIO, with any division "
         "by zero mapped to a missing value. All transformations sit inside a scikit-learn "
         "ColumnTransformer that imputes and standardises numeric columns and imputes and one-hot encodes "
         "categorical columns. The transformer will be fitted on training data only, so no statistic from "
         "the test set can leak into the model. Lastly, the data will be divided with a stratified 80 to 20 "
         "split that preserves the 8.07 percent default rate, giving 246,008 training rows and 61,503 test "
         "rows under a fixed random seed of 42 for reproducibility.")

    h2(doc, "3.3 System Architecture and Contextual Diagram")
    para(doc,
         "Figure 3.2 presents the level 0 data flow diagram, which frames the whole system as a single "
         "process connecting two external entities: the Kaggle data source that supplies the raw CSV files "
         "and the credit analyst who consumes risk information through the dashboard. Internally, the "
         "system decomposes into a small set of Python modules. A data loader reads the raw CSV, fixes the "
         "DAYS_EMPLOYED sentinel and caches the table in parquet format. The preprocessing module applies "
         "the cleaning, feature engineering and encoding plan from Section 3.2. An EDA module produces the "
         "exploratory figures, the model modules will train the classifiers, and an evaluation module will "
         "write every metric into one consolidated metrics file that acts as the single source of truth. "
         "The level 1 decomposition of these processes appears in Appendix B (Figure B.1).")
    add_figure(doc, "dfd_level0",
               "Figure 3.2: Level 0 data flow diagram (contextual diagram) of the loan default analytics system")
    para(doc,
         "The delivery channel is a Streamlit dashboard with three pages. The Overview page, which I "
         "design, shows portfolio KPIs and the key EDA charts computed on a 50,000 row sample so the page "
         "stays responsive. The Model Comparison page, led by my teammate, will read the consolidated "
         "metrics file and display the comparison tables and curves. The Live Prediction page, which I "
         "also design, will collect applicant details through a form, score them with the final model, and "
         "translate the raw score into a Low, Moderate or High risk band using the 50th and 80th "
         "percentiles of sample scores. This banding decision exists because raw probabilities from class "
         "weighted models are inflated, so percentile bands communicate risk more honestly to a non "
         "technical analyst.")

    h2(doc, "3.4 Data Storage Discussion")
    para(doc,
         "The storage design balances simplicity against scale. The 166 MB CSV loads slowly in pandas, so "
         "the data loader caches it as a parquet file, which cuts reload time several fold and makes "
         "repeated experiments practical on a student laptop. For the current single table batch analysis, "
         "flat CSV plus a parquet cache is adequate and keeps the project easy to reproduce. However, the "
         "full relational set is about 2.7 GB across seven tables, and joining bureau and repayment "
         "histories would strain this file based approach. For that extension, a document store such as "
         "MongoDB, which my team and I practised in the module labs, is an attractive option because an "
         "application and its variable length credit history can be stored as one semi structured "
         "document. In a nutshell, the project uses the simplest storage that fits the present scope while "
         "documenting a credible path to a scalable store when the child tables are brought in.")


def build_chapter4(doc):
    h1(doc, "Chapter 4: Model Design and Planning")

    h2(doc, "4.1 Analysis Tools and Techniques")
    para(doc,
         "The technical stack is Python 3.12 with pandas 2.2.2 for data handling, scikit-learn 1.5.0 for "
         "preprocessing and the Logistic Regression and Random Forest models, XGBoost 2.0.3 for my "
         "teammate's boosting model, pyarrow for the parquet cache, matplotlib and seaborn for static EDA "
         "figures, Plotly for interactive charts, and Streamlit 1.36 for the dashboard. Exploration happens "
         "in Jupyter Notebook and the production modules are written in VS Code. Quality is supported by "
         "pytest unit tests covering the loading, preprocessing and evaluation logic. The code lives in a "
         "private GitHub repository with a trunk based workflow and descriptive conventional commit "
         "messages, with the commit work planned evenly between the two members from 15 May to 25 June "
         "2026. The overall schedule, running from 13 April to 28 June 2026 across planning, data "
         "collection, EDA, preprocessing, model development, evaluation, dashboard and final preparation "
         "phases, is shown in the Gantt chart in Appendix A (Figure A.1).")

    h2(doc, "4.2 Hypothesis Testing and Correlation Analysis")
    para(doc,
         "The two hypotheses about attribute relationships were tested on the real data during this "
         "progressive phase. For H2, the chi square test of independence is the appropriate tool because "
         "both the predictor and the target are categorical, and it asks whether the default rate differs "
         "across categories more than chance would allow. Table 4.1 reports the results for the five "
         "socioeconomic attributes named in the hypothesis. Every test rejects independence at far below "
         "the 0.1 percent level, so H2 is supported by the early evidence.")
    add_caption(doc, "Table 4.1: Chi-square tests of independence between categorical attributes and TARGET")
    add_table(doc,
              ["Attribute", "Chi-square", "Degrees of freedom", "p-value"],
              [
                  ("NAME_INCOME_TYPE", "1253.5", "7", "< 0.001"),
                  ("NAME_EDUCATION_TYPE", "1019.2", "4", "< 0.001"),
                  ("CODE_GENDER", "920.8", "2", "< 0.001"),
                  ("NAME_FAMILY_STATUS", "504.7", "5", "< 0.001"),
                  ("NAME_CONTRACT_TYPE", "293.2", "1", "< 0.001"),
              ],
              col_widths=[2.4, 1.3, 1.6, 1.2])
    para(doc,
         "The practical pattern behind these statistics matters to a lender. The default rate falls steadily "
         "with education, from about 11 percent for lower secondary education to the lowest rate among "
         "academic degree holders. Male applicants default at about 10.1 percent against 7.0 percent for "
         "female applicants. Income type separates working applicants from pensioners and students with "
         "visibly different risk levels. Furthermore, the income and credit amount distributions are "
         "heavily right skewed, so the EDA applies a log transform before plotting them, and this skew is "
         "also why median imputation was chosen in Chapter 3.")
    para(doc,
         "For H1, the point biserial correlation measures the linear association between a numeric "
         "attribute and the binary target. Table 4.2 lists the coefficients, all significant at the five "
         "percent level. The three external credit scores show the strongest relationships, with "
         "EXT_SOURCE_3 at r = -0.179, EXT_SOURCE_2 at -0.161 and EXT_SOURCE_1 at -0.155, all negative, "
         "which means higher external scores accompany lower default risk. H1 is therefore supported. Age "
         "and employment length follow at -0.078 and -0.075, confirming that younger and shorter employed "
         "applicants default more often, while raw income is almost uncorrelated with default at -0.004 "
         "once the engineered ratios are available. This ordering mirrors the SHAP ranking reported by "
         "Yang et al. (2025) on the same dataset.")
    add_caption(doc, "Table 4.2: Point-biserial correlations between numeric attributes and TARGET (all p < 0.05)")
    add_table(doc,
              ["Attribute", "r", "Attribute", "r"],
              [
                  ("EXT_SOURCE_3", "-0.179", "AMT_CREDIT", "-0.030"),
                  ("EXT_SOURCE_2", "-0.161", "ANNUITY_INCOME_RATIO", "+0.014"),
                  ("EXT_SOURCE_1", "-0.155", "CREDIT_TERM", "+0.013"),
                  ("AGE_YEARS", "-0.078", "CREDIT_INCOME_RATIO", "-0.008"),
                  ("YEARS_EMPLOYED", "-0.075", "AMT_INCOME_TOTAL", "-0.004"),
              ],
              col_widths=[1.9, 1.0, 2.4, 1.0])
    para(doc,
         "Figure 4.1 visualises the strongest of these relationships by comparing the EXT_SOURCE "
         "distributions for repaid and defaulted loans. The clear separation between the two groups is "
         "exactly what a lender wants from a screening attribute, and it justifies keeping all three "
         "external scores in every model despite their missing values.")
    add_figure(doc, "eda_ext_vs_target",
               "Figure 4.1: Distributions of the external credit scores split by repayment outcome")

    h2(doc, "4.3 Chosen Algorithms and Justification")
    para(doc,
         "The project will train at least two classifiers under my responsibility, selected to answer RQ2. "
         "The first is Logistic Regression configured with the liblinear solver, balanced class weights and "
         "up to 500 iterations. It is the canonical credit scoring baseline (Hand and Henley, 1997), its "
         "coefficients can be explained to a regulator or a declined customer, and the benchmark evidence "
         "shows linear models remain surprisingly competitive on credit data (Baesens et al., 2003). The "
         "balanced class weights make each defaulter count roughly eleven times more during training, which "
         "is the single imbalance correction this model receives.")
    para(doc,
         "The second is a Random Forest with 200 trees, a maximum depth of 20, at least 10 samples per "
         "leaf, and balanced subsample class weighting (Breiman, 2001). The depth and leaf constraints "
         "control overfitting on the wide encoded feature space, while the ensemble captures the non linear "
         "patterns and interactions that the chi square results in Section 4.2 hint at. Its built in "
         "feature importances also feed the analyst facing explanation on the dashboard. Both models will "
         "be compared against the XGBoost model that my teammate is preparing, so the final group report "
         "can test H3 properly, since the literature expects the boosted ensemble to rank best (Lessmann "
         "et al., 2015). Where individual decisions need justification, SHAP style attribution will be "
         "applied to the winning model (Lundberg and Lee, 2017).")

    h2(doc, "4.4 Implementation and Evaluation Plan")
    para(doc,
         "All models will be implemented inside a single scikit-learn Pipeline that shares the one "
         "preprocessor designed in Chapter 3, so every classifier sees identical inputs and the comparison "
         "stays fair. Training will use the stratified 80 to 20 hold-out split, and K-fold cross validation "
         "on the training portion will be used to check that the scores are stable rather than lucky. The "
         "evaluation will report precision, recall, F1 and ROC AUC for the defaulter class together with "
         "full confusion matrices, because accuracy alone is meaningless at an 8 percent default rate. "
         "From the lender perspective, recall measures how many future write-offs the model catches, while "
         "precision measures how many good customers are wrongly flagged, and the right balance depends on "
         "the relative cost of those two errors.")
    para(doc,
         "Two refinements are planned after the first training round. Firstly, a decision threshold sweep "
         "will tune the cut-off away from the default 0.5, since class weighted models shift their score "
         "distributions and the F1 optimal threshold is unlikely to sit at the centre. Secondly, a "
         "calibration check using Brier scores will verify how trustworthy the raw probabilities are. "
         "Because balanced weighting inflates raw scores, the dashboard will deliberately present "
         "percentile based Low, Moderate and High risk bands instead of quoting probabilities, following "
         "the principle that a dashboard must match the decision needs and statistical literacy of its "
         "audience (Sarikaya et al., 2019).")
    para(doc,
         "Lastly, the evaluation plan includes an explicit ethics checkpoint. Gender is statistically "
         "predictive in this data, but using a protected attribute to price or deny credit raises "
         "discrimination concerns, and fairness in credit scoring requires deliberate assessment rather "
         "than silence (Kozodoi, Jacob and Lessmann, 2022). The system is also designed as decision "
         "support, with the analyst making the final call, which keeps it on the right side of the GDPR "
         "restriction on solely automated decisions (European Parliament and Council, 2016; Goodman and "
         "Flaxman, 2017). In a nutshell, the progressive phase has confirmed both data hypotheses, fixed "
         "the preprocessing and architecture design, and set out a fair and measurable plan for the "
         "modelling work that the final group report will complete.")


REFERENCES = [
    "Baesens, B., Van Gestel, T., Viaene, S., Stepanova, M., Suykens, J. and Vanthienen, J. (2003) "
    "'Benchmarking state-of-the-art classification algorithms for credit scoring', Journal of the "
    "Operational Research Society, 54(6), pp. 627-635. doi:10.1057/palgrave.jors.2601545.",

    "Breiman, L. (2001) 'Random forests', Machine Learning, 45(1), pp. 5-32. doi:10.1023/A:1010933404324.",

    "Chawla, N.V., Bowyer, K.W., Hall, L.O. and Kegelmeyer, W.P. (2002) 'SMOTE: Synthetic Minority "
    "Over-sampling Technique', Journal of Artificial Intelligence Research, 16, pp. 321-357. "
    "doi:10.1613/jair.953.",

    "Emmanuel, T., Maupong, T., Mpoeleng, D., Semong, T., Mphago, B. and Tabona, O. (2021) 'A survey on "
    "missing data in machine learning', Journal of Big Data, 8, Article 140. doi:10.1186/s40537-021-00516-9.",

    "European Parliament and Council (2016) Regulation (EU) 2016/679 (General Data Protection Regulation), "
    "Article 22. Official Journal of the European Union, L 119.",

    "Goodman, B. and Flaxman, S. (2017) 'European Union regulations on algorithmic decision-making and a "
    "\"right to explanation\"', AI Magazine, 38(3), pp. 50-57. doi:10.1609/aimag.v38i3.2741.",

    "Hand, D.J. and Henley, W.E. (1997) 'Statistical classification methods in consumer credit scoring: A "
    "review', Journal of the Royal Statistical Society: Series A, 160(3), pp. 523-541. "
    "doi:10.1111/j.1467-985X.1997.00078.x.",

    "Hasan, M.M., Popp, J. and Olah, J. (2020) 'Current landscape and influence of big data on finance', "
    "Journal of Big Data, 7(1), Article 21. doi:10.1186/s40537-020-00291-z.",

    "Kozodoi, N., Jacob, J. and Lessmann, S. (2022) 'Fairness in credit scoring: Assessment, implementation "
    "and profit implications', European Journal of Operational Research, 297(3), pp. 1083-1094. "
    "doi:10.1016/j.ejor.2021.06.023.",

    "Lessmann, S., Baesens, B., Seow, H.-V. and Thomas, L.C. (2015) 'Benchmarking state-of-the-art "
    "classification algorithms for credit scoring: An update of research', European Journal of Operational "
    "Research, 247(1), pp. 124-136. doi:10.1016/j.ejor.2015.05.030.",

    "Lundberg, S.M. and Lee, S.-I. (2017) 'A unified approach to interpreting model predictions', Advances "
    "in Neural Information Processing Systems 30 (NIPS 2017), Long Beach, CA. arXiv:1705.07874.",

    "Sarikaya, A., Correll, M., Bartram, L., Tory, M. and Fisher, D. (2019) 'What do we talk about when we "
    "talk about dashboards?', IEEE Transactions on Visualization and Computer Graphics, 25(1), pp. 682-692. "
    "doi:10.1109/TVCG.2018.2864903.",

    "Srivastava, U. and Gopalkrishnan, S. (2015) 'Impact of big data analytics on banking sector: Learning "
    "for Indian banks', Procedia Computer Science, 50, pp. 643-652. doi:10.1016/j.procs.2015.04.098.",

    "Yang, S., Huang, Z., Xiao, W. and Shen, X. (2025) 'Interpretable credit default prediction with "
    "ensemble learning and SHAP', arXiv preprint arXiv:2505.20815.",
]


def build_references(doc):
    h1(doc, "References")
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(10)


def build_appendices(doc):
    h1(doc, "Appendix A: Project Gantt Chart")
    para(doc,
         "Figure A.1 shows the project schedule from 13 April 2026 to 28 June 2026, covering the planning, "
         "data source identification, data collection plan, EDA, preprocessing, model development, "
         "evaluation, dashboard, testing and deployment, and final preparation phases referenced in "
         "Section 4.1.")
    add_figure(doc, "gantt", "Figure A.1: Project Gantt chart, 13 April 2026 to 28 June 2026")

    page_break(doc)
    h1(doc, "Appendix B: Level 1 Data Flow Diagram")
    para(doc,
         "Figure B.1 decomposes the contextual diagram of Section 3.3 into the individual pipeline "
         "processes: data loading with the parquet cache, preprocessing, exploratory analysis, model "
         "training, evaluation, and the dashboard that serves the credit analyst.")
    add_figure(doc, "dfd_level1", "Figure B.1: Level 1 data flow diagram of the analytics pipeline")

    page_break(doc)
    h1(doc, "Appendix C: Entity Relationship Diagram")
    para(doc,
         "Figure C.1 documents the relational structure of the full Home Credit dataset described in "
         "Section 3.1, with the child tables joining application_train through SK_ID_CURR and "
         "bureau_balance joining bureau through SK_ID_BUREAU. Only application_train is used in the "
         "current modelling scope.")
    add_figure(doc, "erd", "Figure C.1: Entity relationship diagram of the Home Credit tables")


def build_turnitin(doc):
    h1(doc, "Turnitin Similarity Percentage Page")
    para(doc,
         "Placeholder: insert the Turnitin similarity receipt screenshot for this report on this page "
         "before submission.")


# ----------------------------------------------------------------------------
# build + verify
# ----------------------------------------------------------------------------

def build():
    doc = Document()
    setup_styles(doc)

    build_cover(doc)
    page_break(doc)

    para(doc, "Table of Contents", bold=True, size=16)
    add_toc_field(doc)
    page_break(doc)

    build_chapter1(doc)
    page_break(doc)
    build_chapter2(doc)
    page_break(doc)
    build_chapter3(doc)
    page_break(doc)
    build_chapter4(doc)
    page_break(doc)
    build_references(doc)
    page_break(doc)
    build_appendices(doc)
    page_break(doc)
    build_turnitin(doc)

    doc.save(OUT_DOCX)
    print("Saved:", OUT_DOCX)


def verify():
    doc = Document(OUT_DOCX)

    print("\n--- Heading tree ---")
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = int(p.style.name.split()[-1])
            print("  " * (level - 1) + p.text)

    in_body = False
    body_words = 0
    heading_words = 0
    caption_words = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 1":
            if text.startswith("Chapter 1"):
                in_body = True
            elif text.startswith("References"):
                in_body = False
        if not in_body:
            continue
        n = len(re.findall(r"\S+", text))
        if re.match(r"^(Figure|Table) [A-Z0-9]", text):
            caption_words += n
        elif p.style.name.startswith("Heading"):
            heading_words += n
        else:
            body_words += n

    print("\n--- Word counts (Chapters 1-4) ---")
    print("Body prose words:", body_words)
    print("Heading words:", heading_words)
    print("Caption words (excluded):", caption_words)
    print("Prose + headings:", body_words + heading_words)

    n_images = len(doc.inline_shapes)
    print("\nEmbedded inline images:", n_images)
    assert n_images == len(IMG), "expected %d images, found %d" % (len(IMG), n_images)
    for key, path in IMG.items():
        assert os.path.exists(path), "missing image source: " + path
    print("All %d image source paths exist." % len(IMG))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "\u2014" not in full_text, "em-dash found in document"
    print("No em-dashes present.")

    cited = [
        "Hasan, Popp and Olah, 2020", "Srivastava and Gopalkrishnan, 2015",
        "Hand and Henley (1997)", "Baesens et al. (2003)", "Lessmann et al. (2015)",
        "Yang et al. (2025)", "Breiman, 2001", "Chawla et al., 2002",
        "Lundberg and Lee, 2017", "Goodman and Flaxman (2017)",
        "European Parliament and Council, 2016", "Kozodoi, Jacob and Lessmann, 2022",
        "Emmanuel et al. (2021)", "Sarikaya et al., 2019",
    ]
    found = [c for c in cited if c.split("(")[0].split(",")[0].strip() in full_text]
    print("Distinct reference authors found in text: %d" % len(found))


if __name__ == "__main__":
    build()
    verify()
