# -*- coding: ascii -*-
"""Builds the 5011CEM GROUP Final Report docx via python-docx.

Run:  python -X utf8 build_group_report.py
Output: BryanTey_ThongWaiKit_5011CEM_Group_Final_Report.docx (same folder)

All numbers come from 06_Working_Files/REPORT_FACTS_PACK.md (reproduced
2026-06-11). References limited to 06_Working_Files/verified_references.md.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = r"C:\Users\ASUS\Downloads\07_Team_Handoff\07_Team_Handoff"
EDA_DIR = os.path.join(ROOT, "04_Diagrams", "EDA_Plots")
SYS_DIR = os.path.join(ROOT, "04_Diagrams", "System_Diagrams")
FIG_DIR = os.path.join(ROOT, "05_GitHub_and_Code", "code", "reports", "figures")
WORK_DIR = os.path.join(ROOT, "06_Working_Files")
OUT_DIR = os.path.join(ROOT, "08_Final_Submission")
OUT_DOCX = os.path.join(OUT_DIR, "BryanTey_ThongWaiKit_5011CEM_Group_Final_Report.docx")

FONT = "Times New Roman"


# ----------------------------------------------------------------------
# style helpers
# ----------------------------------------------------------------------

def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        srpr = st.element.get_or_add_rPr()
        srfonts = srpr.find(qn("w:rFonts"))
        if srfonts is None:
            srfonts = OxmlElement("w:rFonts")
            srpr.append(srfonts)
        srfonts.set(qn("w:ascii"), FONT)
        srfonts.set(qn("w:hAnsi"), FONT)
        srfonts.set(qn("w:eastAsia"), FONT)

    cap = doc.styles["Caption"]
    cap.font.name = FONT
    cap.font.size = Pt(11)
    cap.font.bold = False
    cap.font.italic = False
    cap.font.color.rgb = RGBColor(0, 0, 0)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_after = Pt(12)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


# ----------------------------------------------------------------------
# content helpers
# ----------------------------------------------------------------------

def para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False,
         italic=False, center=False, space_after=None):
    p = doc.add_paragraph()
    if center:
        align = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 12}[level], bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def figure(doc, path, cap, width=6.0):
    assert os.path.exists(path), "MISSING IMAGE: " + path
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def table(doc, header, rows, cap=None, col_widths=None, font_size=11):
    if cap:
        caption(doc, cap)
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    all_rows = [header] + rows
    for r, row in enumerate(all_rows):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.line_spacing = 1.0
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(str(val))
            set_run_font(run, size=font_size, bold=(r == 0))
    if col_widths:
        for c, w in enumerate(col_widths):
            for r in range(len(all_rows)):
                t.cell(r, c).width = Inches(w)
    return t


def kv_table(doc, pairs, label_width=2.4, value_width=4.0):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, (k, v) in enumerate(pairs):
        for c, val in enumerate((k, v)):
            cell = t.cell(r, c)
            cp = cell.paragraphs[0]
            cp.paragraph_format.line_spacing = 1.15
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(val)
            set_run_font(run, size=11, bold=(c == 0))
        t.cell(r, 0).width = Inches(label_width)
        t.cell(r, 1).width = Inches(value_width)
    return t


def toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    set_run_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Right-click this line and choose Update Field to generate the Table of Contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, txt, fld_end):
        run._element.append(el)


def page_break(doc):
    doc.add_page_break()


# ----------------------------------------------------------------------
# document sections
# ----------------------------------------------------------------------

def build_cover(doc):
    para(doc, "INTI International College Penang", center=True, bold=True, size=14, space_after=0)
    para(doc, "School of Computing", center=True, bold=True, size=14, space_after=0)
    para(doc, "3+0 Bachelor of Science (Hons) in Computer Science, "
              "in collaboration with Coventry University, UK", center=True, size=12, space_after=6)
    para(doc, "Coursework cover sheet", center=True, bold=True, size=13, space_after=6)
    para(doc, "Section A - To be completed by the student", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    kv_table(doc, [
        ("Full Name", "Bryan Tey Kai Yuan and Thong Wai Kit"),
        ("CU Student ID Number", "P23015693 (Bryan Tey Kai Yuan) / P23015668 (Thong Wai Kit)"),
        ("Semester", "April 2026"),
        ("Lecturer", "Ms. Vimala Doraisamy"),
        ("Module Code and Title", "5011CEM Big Data Programming Project"),
        ("Assignment No. / Title", "Home Credit Loan-Default Prediction - 100% of Module Mark "
                                   "(Group Final Report)"),
        ("Hand out date", "Week 1"),
        ("Due date", "Individual Progress Report - WK10 / Final Report - WK12 / VIVA - WK13 & 14"),
    ])
    para(doc, "")
    para(doc, "Penalties: No late work will be accepted. If you are unable to submit coursework "
              "on time due to extenuating circumstances, you may be eligible for an extension. "
              "Please consult the lecturer.", size=11)
    para(doc, "Declaration: I/we the undersigned confirm that I/we have read and agree to abide "
              "by the University regulations on plagiarism and cheating and Faculty coursework "
              "policies and procedures. I/we confirm that this piece of work is my/our own. "
              "I/we consent to appropriate storage of our work for plagiarism checking.", size=11)
    para(doc, "Signature(s): Bryan Tey Kai Yuan, Thong Wai Kit", size=11, space_after=18)
    para(doc, "Group Final Report", center=True, bold=True, size=16, space_after=0)
    para(doc, "Home Credit Loan-Default Prediction", center=True, bold=True, size=14, space_after=0)
    para(doc, "Banking and Financial Sector Scenario", center=True, size=12, space_after=0)


def build_toc(doc):
    para(doc, "Table of Contents", center=True, bold=True, size=14, space_after=6)
    toc_field(doc)


def build_ch1(doc):
    heading(doc, "Chapter 1: Descriptive Statistics and Exploratory Data Analysis", 1)
    para(doc, "This chapter presents the dataset our team analysed, the types of attributes it "
              "contains, the distributions and group differences we observed, and the formal "
              "statistical tests that justified our variable selection. Every chart referred to "
              "here was generated programmatically by our EDA module, so the figures can be "
              "regenerated at any time from the raw data.")

    heading(doc, "1.1 Dataset Overview and Attribute Types", 2)
    para(doc, "Our team built this project on the Home Credit Default Risk dataset, which Home "
              "Credit Group released for a public Kaggle competition in 2018. The core table, "
              "application_train.csv, holds 307,511 loan applications described by 122 columns, "
              "and the raw file takes up about 166 MB on disk. Each row represents one consumer "
              "loan application. The label TARGET equals 1 when the applicant experienced payment "
              "difficulty and 0 when the loan was repaid without serious problems. Out of the "
              "307,511 applications, 24,825 are defaults, which gives a default rate of 8.07 "
              "percent. The data is real, anonymised and publicly released, it carries no direct "
              "identifiers, and it comfortably exceeds the module requirement of at least 100,000 "
              "records.")
    para(doc, "The application table does not stand alone. Six relational child tables, namely "
              "bureau (with bureau_balance underneath it), previous_application, "
              "POS_CASH_balance, installments_payments and credit_card_balance, join the main "
              "table on the applicant key SK_ID_CURR and together span about 2.7 GB. After "
              "studying the supplied data dictionary, HomeCredit_columns_description.csv, we "
              "decided to scope the deployed model to the single application table and to "
              "document the child tables as a future extension, since the main table already "
              "carries the demographic, financial and external score attributes that the "
              "literature identifies as the core of credit scoring.")
    para(doc, "The 122 attributes fall into two broad families. Firstly, there are categorical "
              "attributes that store applicant circumstances as labels. Examples include "
              "NAME_CONTRACT_TYPE (Cash loans or Revolving loans), CODE_GENDER, "
              "NAME_EDUCATION_TYPE, NAME_INCOME_TYPE, NAME_FAMILY_STATUS and OCCUPATION_TYPE. "
              "Secondly, there are numerical attributes such as AMT_INCOME_TOTAL, AMT_CREDIT and "
              "AMT_ANNUITY for the financial amounts, the day counters DAYS_BIRTH and "
              "DAYS_EMPLOYED for age and employment history, and the three normalised external "
              "credit scores EXT_SOURCE_1, EXT_SOURCE_2 and EXT_SOURCE_3. In addition, the table "
              "carries a block of binary FLAG_DOCUMENT indicators that record which supporting "
              "documents an applicant submitted. Recognising these data types early mattered "
              "because our preprocessing pipeline routes numeric and categorical columns through "
              "different transformations, as Chapter 2 explains.")
    para(doc, "Data quality was a genuine concern. Roughly 24 percent of all cells in the table "
              "are missing, and the missingness concentrates in a sparse building survey block "
              "(COMMONAREA, NONLIVINGAPARTMENTS, YEARS_BUILD and similar columns) that is mostly "
              "empty, as Figure A.5 in Appendix A shows. Besides that, the DAYS_EMPLOYED column "
              "contains the sentinel value 365243, which corresponds to roughly one thousand "
              "years of employment and actually marks pensioners and unemployed applicants. This "
              "sentinel affects 55,374 rows, or 18.0 percent of the data, and we replaced it with "
              "a missing value at load time so it could not distort the employment features. "
              "Guided by the survey on missing data in machine learning by Emmanuel et al. "
              "(2021), we chose imputation over row deletion, because deleting incomplete rows "
              "would have discarded a large share of the dataset and risked biasing the sample "
              "that remained.")

    heading(doc, "1.2 Class Balance of the Target Variable", 2)
    para(doc, "Figure 1.1 shows the class balance of the target. Repaid loans make up 91.9 "
              "percent of applications while defaults make up only 8.1 percent. This imbalance "
              "shaped almost every later decision in the project. With barely 8 percent "
              "positives, plain accuracy becomes misleading, because a model that rejects nobody "
              "still scores about 92 percent accuracy while catching zero defaulters. He and "
              "Garcia (2009) describe this trap in detail and recommend judging classifiers "
              "through precision, recall, F1 and ROC analysis instead, which is exactly the "
              "metric set we adopted. Moreover, the imbalance motivated the class weighting and "
              "scale_pos_weight corrections built into the models in Chapter 2.")
    figure(doc, os.path.join(EDA_DIR, "eda_target.png"),
           "Figure 1.1: Class balance of TARGET (91.9 percent repaid against 8.1 percent default).")

    heading(doc, "1.3 Distributions and Group Differences", 2)
    para(doc, "The financial amount columns are heavily right-skewed. Most applicants earn modest "
              "incomes and borrow moderate amounts, while a small tail of very large incomes and "
              "credits stretches the axis. A logarithmic transform applied during the exploratory "
              "work brought the income distribution close to symmetric, which confirmed that the "
              "skew comes from scale rather than from data entry errors. The raw distribution "
              "plots are provided in Appendix A as Figure A.1, and the comparison of credit and "
              "income amounts across the two outcome classes appears as Figure A.2.")
    para(doc, "Figure 1.2 compares the default rate across categorical groups, and the pattern is "
              "consistent and economically sensible. The default rate falls steadily as education "
              "level rises, from about 11 percent for applicants with lower secondary education "
              "down to the lowest rate for the small academic degree group. Male applicants "
              "default more often (about 10.1 percent) than female applicants (about 7.0 "
              "percent), and income type separates the groups clearly as well. These differences "
              "guided our choice of categorical features, and they also raised an ethical "
              "question about demographic attributes that we return to in Chapter 4.")
    figure(doc, os.path.join(EDA_DIR, "eda_default_by_cat.png"),
           "Figure 1.2: Default rate by education level, gender, income type and family status.")
    para(doc, "The three external credit scores turned out to be the most informative single "
              "attributes in the table. Figure 1.3 plots their values for repaid and defaulted "
              "loans, and the separation is visible by eye, since defaulters sit clearly lower on "
              "EXT_SOURCE_2 and EXT_SOURCE_3. Age and employment length show a similar but weaker "
              "pattern, with younger and recently employed applicants defaulting more often, as "
              "Figure A.4 in Appendix A illustrates. The standalone distributions of the three "
              "external scores appear in Figure A.3.")
    figure(doc, os.path.join(EDA_DIR, "eda_ext_vs_target.png"),
           "Figure 1.3: External credit scores split by loan outcome.")

    heading(doc, "1.4 Statistical Tests and Variable Selection", 2)
    para(doc, "We did not select variables by eye alone. For the categorical attributes we ran "
              "chi-square tests of independence against TARGET, and every headline attribute was "
              "significant at p < 0.001. NAME_INCOME_TYPE produced the largest statistic "
              "(chi-square = 1253.5 with 7 degrees of freedom), followed by NAME_EDUCATION_TYPE "
              "(1019.2, dof 4), CODE_GENDER (920.8, dof 2), NAME_FAMILY_STATUS (504.7, dof 5) and "
              "NAME_CONTRACT_TYPE (293.2, dof 1). These results confirmed that the group "
              "differences seen in Figure 1.2 are systematic rather than sampling noise.")
    para(doc, "For the numeric attributes we computed point-biserial correlations with TARGET, "
              "all significant at p < 0.05. The external scores dominate, with EXT_SOURCE_3 at "
              "r = -0.179, EXT_SOURCE_2 at r = -0.161 and EXT_SOURCE_1 at r = -0.155. AGE_YEARS "
              "(r = -0.078) and YEARS_EMPLOYED (r = -0.075) follow, while AMT_CREDIT (r = -0.030) "
              "and raw income (r = -0.004) barely matter on their own. Two engineered ratios, "
              "ANNUITY_INCOME_RATIO (r = +0.014) and CREDIT_TERM (r = +0.013), correlate "
              "positively with default, which fits intuition because heavier repayment burdens "
              "are riskier. The full correlation heatmap is given as Figure A.6 in Appendix A. "
              "The dominance of the external scores agrees with Yang et al. (2025), who found "
              "that EXT_SOURCE features dominate SHAP attributions on this same dataset.")
    para(doc, "Based on this evidence, we kept the informative numeric and categorical "
              "attributes, dropped every column with 60 percent or more missing values, dropped "
              "the weak FLAG_DOCUMENT indicators, and engineered six domain ratios (described in "
              "Chapter 2) to expose repayment burden directly. Raw income was retained mainly "
              "because the ratios derived from it carry the real signal.")


def build_ch2(doc):
    heading(doc, "Chapter 2: Model Development", 1)
    para(doc, "This chapter describes the functionality of every code module in the project, what "
              "each file is for and how it achieves its job, including testing. We stay at the "
              "design level throughout and avoid syntax detail. The technology stack was Python "
              "3.12 with pandas 2.2.2 for data handling, scikit-learn 1.5.0 for the preprocessing "
              "pipeline and the baseline models, XGBoost 2.0.3 for gradient boosting, "
              "imbalanced-learn for the SMOTE experiment, pyarrow for the parquet cache, "
              "matplotlib and seaborn for static figures, Plotly and Streamlit 1.36 for the "
              "dashboard, joblib for model persistence and pytest for automated testing. "
              "Development was done in Jupyter Notebook for exploration and VS Code for the "
              "package code, so quick experiments and production modules stayed cleanly "
              "separated.")

    heading(doc, "2.1 Pipeline Architecture", 2)
    para(doc, "The implementation is a Python 3.12 package organised under a src folder, with "
              "separate modules for data loading, preprocessing, exploratory analysis, models, "
              "evaluation and the dashboard. The context diagram (Figure C.1) and the level 1 "
              "data flow diagram (Figure C.2) in Appendix C show how the raw Kaggle data flows "
              "through cleaning, feature engineering and model training into figures, a "
              "consolidated metrics file and the Streamlit dashboard. The entity relationship "
              "diagram (Figure C.3) documents how the six relational child tables join "
              "application_train on SK_ID_CURR, and the UML class diagram (Figure C.4) records "
              "the module structure. One deliberate design rule was that the pipeline writes a "
              "single consolidated reports/metrics.csv file and the dashboard reads that same "
              "file, so the numbers shown on screen can never drift from the numbers in this "
              "report.")

    heading(doc, "2.2 Data Loading (data_loader)", 2)
    para(doc, "The data_loader module is responsible for getting the 166 MB CSV into memory "
              "quickly and consistently. On the first run it parses application_train.csv with "
              "pandas and saves a parquet copy through pyarrow; on every later run it loads the "
              "parquet cache instead, which is several times faster than re-parsing the CSV. The "
              "loader is also where the DAYS_EMPLOYED sentinel from Chapter 1 is converted to a "
              "missing value, so every downstream module receives data with the quirk already "
              "fixed. Centralising the fix in one place prevented the classic mistake where one "
              "module remembers a correction and another forgets it.")

    heading(doc, "2.3 Preprocessing (preprocessing)", 2)
    para(doc, "The preprocessing module turns the raw table into model-ready matrices in five "
              "stages, summarised by the flowchart in Figure C.5. Firstly, clean_data removes "
              "exact duplicate rows (none existed in the real data) and logs the missing-value "
              "footprint instead of deleting incomplete rows. Secondly, engineer_features adds "
              "six domain ratios: AGE_YEARS, YEARS_EMPLOYED, CREDIT_INCOME_RATIO, "
              "ANNUITY_INCOME_RATIO, CREDIT_TERM and EMPLOYED_AGE_RATIO, with any division by "
              "zero converted to a missing value. Thirdly, select_feature_columns drops the "
              "identifier and the target, the columns with 60 percent or more missing values and "
              "the FLAG_DOCUMENT indicators. Fourthly, a ColumnTransformer routes numeric columns "
              "through median imputation and standard scaling, and categorical columns through "
              "mode imputation and one-hot encoding that ignores categories never seen in "
              "training. The transformer is fitted on training data only, so no information leaks "
              "from the test set into the model. Lastly, stratified_split produces an 80/20 "
              "hold-out that preserves the 8.07 percent default rate, giving 246,008 training "
              "rows and 61,503 test rows under a fixed random seed for reproducibility. This "
              "staged design follows the big data preprocessing guidance of Garcia et al. (2016), "
              "which treats cleaning and feature selection as first-class pipeline stages rather "
              "than ad hoc scripts.")

    heading(doc, "2.4 Exploratory Analysis (eda)", 2)
    para(doc, "The eda module generates every descriptive chart used in Chapter 1 and Appendix A "
              "directly from the loaded data, including the class balance chart, the default rate "
              "breakdowns, the distribution plots and the correlation heatmap. Keeping chart "
              "generation in code rather than in a throwaway notebook means the figures "
              "regenerate identically whenever the pipeline reruns, which protected the report "
              "from stale graphics.")

    heading(doc, "2.5 Model Modules (models)", 2)
    para(doc, "Three classifiers live in the models package, and each one is wrapped in a "
              "scikit-learn Pipeline together with the shared preprocessor, so a saved model "
              "always carries its own preprocessing. Logistic regression served as our "
              "interpretable baseline, in line with its long-standing role as the industry "
              "standard scorecard method (Hand and Henley, 1997); we used the liblinear solver "
              "with balanced class weights and up to 500 iterations. The random forest (Breiman, "
              "2001) used 200 trees with maximum depth 20, at least 10 samples per leaf and "
              "balanced subsample class weights to cope with the skewed classes. XGBoost (Chen "
              "and Guestrin, 2016) was configured with 400 trees, maximum depth 6, learning rate "
              "0.1, row and column subsampling of 0.9 and the histogram tree method. For the "
              "imbalance we set scale_pos_weight to the negative-to-positive ratio of the "
              "training fold, about 11.4, following the official tuning guidance (XGBoost "
              "Developers, 2022), and we monitored the area under the precision-recall curve "
              "during training because it reacts more sharply than ROC-AUC at this imbalance "
              "level.")

    heading(doc, "2.6 Evaluation (evaluation)", 2)
    para(doc, "The evaluation module computes accuracy, precision, recall, F1 and ROC-AUC for any "
              "fitted pipeline on the hold-out set, renders the confusion matrix and ROC curve "
              "figures, and appends each model's results to the consolidated metrics.csv. "
              "Because every model passes through the same evaluation function on the same "
              "61,503 hold-out applications, the comparison in Section 2.10 is guaranteed to be "
              "like for like.")

    heading(doc, "2.7 Orchestration and Experiment Scripts", 2)
    para(doc, "Three scripts drive the package. run_pipeline.py is the single entry point: it "
              "produces the EDA figures, trains the three models, evaluates them and writes the "
              "consolidated metrics file in one command. xgb_experiment.py holds our controlled "
              "class-imbalance experiment. On the same hold-out it compared a baseline XGBoost "
              "with no correction, the scale_pos_weight correction alone, SMOTE oversampling "
              "(Chawla et al., 2002) to minority ratios of 0.3 and 0.5, and combinations of "
              "SMOTE with scale_pos_weight. The scale_pos_weight-only configuration achieved the "
              "best F1 at the default threshold, while combining SMOTE with scale_pos_weight "
              "double-corrected the imbalance, pushing recall up but crushing precision. This "
              "experiment is the reason the deployed model uses a single correction only. "
              "Lastly, optimise_models.py performs the optimisation work reported in Chapter 3: "
              "it sweeps the decision threshold, draws calibration curves, measures training and "
              "prediction runtimes and extracts feature importances.")

    heading(doc, "2.8 Dashboard (dashboard/app)", 2)
    para(doc, "The Streamlit dashboard presents the results to a non-technical audience across "
              "three pages. The Overview page shows portfolio KPIs and the key EDA charts "
              "computed on a 50,000-row sample so the page stays responsive. The Model Comparison "
              "page reads metrics.csv and displays the metrics table, the threshold tuning table, "
              "the confusion matrices, the ROC curves and the calibration and importance figures. "
              "The Live Prediction page collects applicant details through a form, scores them "
              "with the trained XGBoost pipeline and reports a Low, Moderate or High risk band "
              "based on the 50th and 80th percentiles of sample scores rather than a raw "
              "probability, for the calibration reasons explained in Chapter 3. We deliberately "
              "separated the audience-facing summary from the analyst-facing comparison, "
              "following the dashboard design dimensions identified by Sarikaya et al. (2019).")

    heading(doc, "2.9 Automated Tests (tests)", 2)
    para(doc, "The tests package holds 20 pytest unit tests covering the loader, the cleaning and "
              "feature engineering functions, the feature selection rules, the stratified split "
              "proportions and the evaluation metrics, plus a Streamlit AppTest smoke test that "
              "boots every dashboard page and asserts that it renders without errors. All 20 "
              "tests passed in the final run. How testing fitted into our day-to-day workflow is "
              "discussed in Chapter 4.")

    heading(doc, "2.10 Hold-out Results", 2)
    para(doc, "Table 2.1 reports the hold-out performance of the three models on the 61,503 "
              "unseen applications.")
    table(doc,
          ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
          [
              ["XGBoost", "0.759", "0.190", "0.612", "0.290", "0.763"],
              ["Random Forest", "0.861", "0.240", "0.332", "0.278", "0.742"],
              ["Logistic Regression", "0.691", "0.162", "0.676", "0.261", "0.749"],
          ],
          cap="Table 2.1: Hold-out performance of the three models (61,503 applications).",
          col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9])
    para(doc, "XGBoost was the strongest model overall and is the one deployed behind the "
              "dashboard. It achieved the best F1 (0.290) and the best ROC-AUC (0.763), catching "
              "61.2 percent of true defaulters at 19.0 percent precision. Its confusion matrix "
              "(Figure 2.1) shows 3,037 of the 4,965 defaulters correctly flagged at the cost of "
              "12,924 false alarms, and its ROC curve appears in Figure 2.2. The random forest "
              "looked the most accurate at 0.861, but the accuracy is hollow because it missed "
              "about two thirds of the defaulters (recall 0.332, only 1,646 caught). Logistic "
              "regression sat at the other extreme, catching 67.6 percent of defaulters but "
              "wrongly flagging 17,417 good customers. The confusion matrices and ROC curves for "
              "these two models are provided in Appendix B as Figures B.1 to B.4. The overall "
              "ranking, with the boosted and bagged ensembles ahead of the linear baseline, "
              "matches the large credit scoring benchmark of Lessmann et al. (2015).")
    figure(doc, os.path.join(FIG_DIR, "cm_xgboost.png"),
           "Figure 2.1: Confusion matrix of XGBoost on the 61,503-application hold-out.")
    figure(doc, os.path.join(FIG_DIR, "roc_xgboost.png"),
           "Figure 2.2: ROC curve of XGBoost (AUC = 0.763).")
    para(doc, "Figure 2.3 shows the feature importances. The three external credit scores "
              "dominate for both ensemble models, followed by CREDIT_TERM, age and employment "
              "length, while for XGBoost the higher education indicator, gender and the revolving "
              "loan contract type also contribute. Attribution of this kind is the practical "
              "route to explaining individual predictions (Lundberg and Lee, 2017), and the "
              "pattern we observed is consistent with the SHAP analysis of Yang et al. (2025) on "
              "the same dataset. One caution emerged from the logistic regression coefficients: "
              "its largest absolute weights belong to rare categories such as the Academic "
              "degree education level and the Student and Pensioner income types, which contain "
              "very few applicants. This is a known weakness of linear models on one-hot encoded "
              "data, where thin categories receive extreme weights, and it reinforced our "
              "preference for the tree ensembles whose importances concentrated on broadly "
              "populated features instead.")
    figure(doc, os.path.join(FIG_DIR, "feature_importance.png"),
           "Figure 2.3: Feature importances of the trained models.")


def build_ch3(doc):
    heading(doc, "Chapter 3: Measuring and Optimising Algorithm Complexity", 1)
    para(doc, "Big data programming is not only about accuracy; the cost of computing also "
              "matters once a pipeline reruns many times a day. In this chapter we compare the "
              "three algorithms on theoretical complexity and on measured wall-clock time, then "
              "present the two optimisations that improved results without retraining "
              "(threshold tuning) and the two that improved speed (parquet caching and the "
              "histogram tree method).")

    heading(doc, "3.1 Theoretical Complexity and Measured Runtimes", 2)
    para(doc, "We assessed each algorithm twice, firstly on paper through asymptotic complexity "
              "and secondly with wall-clock measurements on the actual 246,008-row training set "
              "and 61,503-row test set. Table 3.1 places both views side by side.")
    table(doc,
          ["Model", "Training complexity", "Fit time (s)", "Predict time (s)"],
          [
              ["Logistic Regression", "O(n x d) per epoch (liblinear, single-threaded)", "95.0", "0.5"],
              ["Random Forest", "O(T x n log n x sqrt(d)), T = 200", "58.8", "1.2"],
              ["XGBoost", "O(T x L x n) with histogram binning, T = 400", "22.8", "0.9"],
          ],
          cap="Table 3.1: Theoretical training complexity and measured runtimes "
              "(246,008 training rows, 61,503 test rows).",
          col_widths=[1.5, 3.0, 1.0, 1.0])
    para(doc, "The measurements tell a more interesting story than the formulas. XGBoost trained "
              "fastest at 22.8 seconds even though it builds the most trees, because the "
              "histogram method bins every feature once and then evaluates splits over bins "
              "instead of raw values, and the work parallelises across cores (Chen and Guestrin, "
              "2016). Logistic regression was slowest at 95.0 seconds despite being the simplest "
              "model, because the liblinear solver is single-threaded and the one-hot encoding "
              "inflates the design matrix to roughly 240 columns. The random forest sat in "
              "between at 58.8 seconds. Prediction is fast for all three, between 0.5 and 1.2 "
              "seconds for the full test set, so any of the models could score applications "
              "interactively.")

    heading(doc, "3.2 Decision Threshold Tuning", 2)
    para(doc, "The cheapest optimisation we found did not retrain anything. All three classifiers "
              "output scores, and the default decision threshold of 0.5 is arbitrary once class "
              "weighting shifts the score distribution. optimise_models.py therefore swept the "
              "threshold across its range and selected the F1-maximising point for each model, "
              "as Figure 3.1 shows and Table 3.2 summarises.")
    table(doc,
          ["Model", "F1 at 0.5", "Best threshold", "F1 at best", "Precision change", "Recall change"],
          [
              ["Logistic Regression", "0.261", "0.66", "0.300", "0.162 -> 0.234", "0.676 -> 0.416"],
              ["Random Forest", "0.278", "0.47", "0.285", "0.240 -> 0.224", "0.332 -> 0.393"],
              ["XGBoost", "0.290", "0.65", "0.315", "0.190 -> 0.260", "0.612 -> 0.400"],
          ],
          cap="Table 3.2: Decision threshold tuning on the hold-out set.",
          col_widths=[1.6, 0.9, 1.0, 0.9, 1.3, 1.3])
    figure(doc, os.path.join(FIG_DIR, "threshold_sweep.png"),
           "Figure 3.1: Decision threshold sweep with the F1-maximising point marked per model.")
    para(doc, "For XGBoost the tuned threshold of 0.65 lifted F1 from 0.290 to 0.315 and raised "
              "precision from 0.190 to 0.260, in exchange for recall falling from 0.612 to "
              "0.400. Logistic regression gained the most in relative terms, moving from 0.261 "
              "to 0.300 at a threshold of 0.66. In a real lending business, however, the "
              "operating threshold should come from the cost ratio between a missed default and "
              "a wrongly rejected customer rather than from F1 alone, and we return to this "
              "point in Chapter 5.")

    heading(doc, "3.3 Probability Calibration", 2)
    para(doc, "Because the imbalance corrections inflate the raw scores, we checked calibration "
              "explicitly. The Brier scores were 0.124 for the random forest, 0.163 for XGBoost "
              "and 0.202 for logistic regression, and all three calibration curves sit above the "
              "diagonal (Figure 3.2), which means every model systematically overstates the "
              "default probability. We therefore treat the outputs as ranking scores rather than "
              "probabilities, and the dashboard maps them to Low, Moderate and High percentile "
              "bands instead of quoting raw numbers to the user.")
    figure(doc, os.path.join(FIG_DIR, "calibration_curves.png"),
           "Figure 3.2: Calibration curves of the three models (Brier scores: RF 0.124, "
           "XGBoost 0.163, LR 0.202).")

    heading(doc, "3.4 Data and Storage Optimisations", 2)
    para(doc, "Two further optimisations target the data layer rather than the models. Firstly, "
              "the parquet cache in data_loader cuts reload time several-fold compared with "
              "parsing the 166 MB CSV, which mattered because the pipeline was rerun dozens of "
              "times during development. Secondly, we considered storage beyond the single "
              "table. The full relational dataset spans about 2.7 GB across seven tables, and at "
              "that scale a flat CSV workflow becomes clumsy. A document store such as MongoDB, "
              "which we practised with during the labs, would suit the semi-structured "
              "application records well. On the other hand, for the current single-table batch "
              "analysis the CSV plus parquet arrangement remained adequate and far simpler to "
              "reproduce, so we kept it and documented the migration path instead. The parquet "
              "gain comes from its columnar layout and typed binary encoding, which avoids the "
              "row-by-row text parsing and type inference that make CSV loading slow, and the "
              "same columnar format would remain the natural interchange format if the project "
              "later moved to a distributed engine.")


def build_ch4(doc):
    heading(doc, "Chapter 4: Professional Practices", 1)

    heading(doc, "4.1 Version Control and Collaboration", 2)
    para(doc, "All code lived in a private GitHub repository "
              "(gokuthong/5011cem-paysim-fraud-detection) shared between the two members. We "
              "worked trunk-based on the main branch with small, frequent commits, 36 commits in "
              "total split evenly at 18 per member between 15 May and 25 June 2026. Commit "
              "messages followed a descriptive conventional style, with feat, fix, docs and test "
              "prefixes followed by a short imperative summary, which kept the history readable "
              "and made it easy to trace which change introduced which figure or metric. Because "
              "the team had only two members, we adopted dual roles instead of formal pull "
              "requests: each member reviewed the other's commits after every push, and anything "
              "unclear was discussed before the next piece of work started. The schedule this "
              "workflow followed is recorded in the Gantt chart in Appendix D (Figure D.1), "
              "covering the full project window from 13 April to 28 June 2026.")

    heading(doc, "4.2 Testing Discipline", 2)
    para(doc, "Testing was continuous rather than a final phase. The 20 pytest unit tests "
              "described in Chapter 2 ran before every significant commit, which caught "
              "regressions early, for example when a change to the feature selection rules "
              "altered the column count that the model tests expected. The Streamlit AppTest "
              "smoke test gave us confidence that dashboard refactoring never silently broke a "
              "page, since it boots each page headlessly and fails on any rendering error. "
              "Keeping the whole suite fast was a deliberate choice, so running it never felt "
              "like a burden and therefore actually happened.")

    heading(doc, "4.3 Ethical and Legal Considerations", 2)
    para(doc, "Credit scoring sits in ethically sensitive territory, so we handled both the data "
              "and the model's role with care. The dataset is anonymised and was released openly "
              "by Home Credit Group, so no personal identifiers were processed at any point. The "
              "harder questions concern deployment. Under Article 22 of the GDPR, individuals "
              "have the right not to be subject to a decision based solely on automated "
              "processing that significantly affects them (European Parliament and Council, "
              "2016), and a loan rejection is exactly such a decision. The associated right to "
              "explanation (Goodman and Flaxman, 2017) means a deployed scorer must be able to "
              "justify individual outcomes, which is why we kept an interpretable baseline and "
              "feature attributions alongside the boosted model. Furthermore, our EDA showed "
              "gender to be predictive of default, but using a protected attribute in credit "
              "decisions raises discrimination concerns, and fairness research shows that "
              "accuracy and fairness must be balanced deliberately rather than assumed (Kozodoi "
              "et al., 2022). We therefore flag the demographic features as an ethical "
              "limitation rather than a modelling success, and the system is positioned as "
              "decision support for a human credit officer, not as an automatic decision maker.")


def build_ch5(doc):
    heading(doc, "Chapter 5: Conclusions and Recommendations", 1)

    heading(doc, "5.1 Findings Against Goals and Hypotheses", 2)
    para(doc, "The project set out to predict loan default from the single Home Credit "
              "application table, to compare an interpretable baseline against ensemble methods "
              "under severe class imbalance, and to serve the results through an interactive "
              "dashboard. All three goals were met. Our first hypothesis, that boosted ensembles "
              "would beat the linear baseline under imbalance, was supported: XGBoost achieved "
              "the best F1 (0.290, rising to 0.315 after threshold tuning) and the best ROC-AUC "
              "(0.763), against 0.261 and 0.749 for logistic regression, in line with published "
              "benchmarks (Lessmann et al., 2015). Our second hypothesis, that the external "
              "credit scores would be the strongest predictors, was supported by both the "
              "point-biserial correlations, where EXT_SOURCE_3 at r = -0.179 led every other "
              "numeric attribute, and by the model importances, matching independent findings on "
              "this dataset (Yang et al., 2025). Thirdly, the imbalance experiment showed that a "
              "single correction through scale_pos_weight outperformed every SMOTE combination, "
              "confirming that stacking corrections double-counts the minority class.")

    heading(doc, "5.2 Limitations", 2)
    para(doc, "Firstly, the deployed model uses only application_train; the six relational child "
              "tables, which hold bureau records and repayment histories, remain untapped, and "
              "competition experience suggests they carry substantial extra signal. Secondly, "
              "the models are poorly calibrated, with all three curves above the diagonal, so "
              "the raw outputs cannot be quoted as probabilities. Thirdly, precision remains "
              "modest even after tuning (0.260 for XGBoost), which means roughly three out of "
              "every four flagged applicants would actually have repaid; the model is a "
              "screening tool, not a verdict. Fourthly, demographic attributes such as gender "
              "contribute predictive power, but their use is ethically and legally questionable, "
              "as Chapter 4 discussed. Lastly, the runtimes were measured on a single local "
              "machine, so the complexity comparison is indicative rather than general.")

    heading(doc, "5.3 Recommendations", 2)
    para(doc, "We recommend five concrete next steps. Firstly, the operating threshold should be "
              "set from the lender's actual cost matrix, weighing the loss from a missed default "
              "against the margin lost on a wrongly rejected customer, instead of maximising F1. "
              "Secondly, an explicit calibration stage, Platt scaling or isotonic regression "
              "fitted on a separate validation fold, should be added before any probability is "
              "shown to users. Thirdly, feature engineering should be extended to the relational "
              "child tables, starting with bureau credit history and previous application "
              "outcomes. Fourthly, a deployed model needs monitoring: the score distribution and "
              "the realised default rate should be tracked monthly and the model retrained when "
              "drift appears. Lastly, a fairness audit should be carried out before any "
              "production use, measuring group-wise error rates and testing the effect of "
              "removing protected attributes (Kozodoi et al., 2022). In a nutshell, the project "
              "delivered a working, tested and explainable default-screening pipeline, and the "
              "path from screening tool to production credit model is now clearly mapped.")


def build_ch6(doc):
    heading(doc, "Chapter 6: Division of Responsibilities", 1)
    para(doc, "Both members contributed throughout the project, but every area had one clear "
              "owner so that nothing fell between two stools. Table 6.1 records the division. "
              "Bryan led the exploratory analysis and the two interpretable models, built the "
              "Overview and Live Prediction dashboard pages and assembled this report. Thong led "
              "the data engineering, the boosted model and its imbalance experiment, the "
              "evaluation module, the Model Comparison page and the architecture documentation. "
              "Code review was mutual, as described in Chapter 4.")
    table(doc,
          ["Project Area", "Lead Member"],
          [
              ["Exploratory data analysis and statistical tests", "Bryan Tey Kai Yuan"],
              ["Logistic Regression and Random Forest models", "Bryan Tey Kai Yuan"],
              ["Dashboard: Overview and Live Prediction pages", "Bryan Tey Kai Yuan"],
              ["Group report assembly", "Bryan Tey Kai Yuan"],
              ["Data loading and preprocessing modules", "Thong Wai Kit"],
              ["XGBoost model and class-imbalance experiment", "Thong Wai Kit"],
              ["Evaluation module and metrics consolidation", "Thong Wai Kit"],
              ["Dashboard: Model Comparison page", "Thong Wai Kit"],
              ["Architecture documentation and system diagrams", "Thong Wai Kit"],
              ["Code review, testing and viva preparation", "Both members"],
          ],
          cap="Table 6.1: Division of responsibilities between the two members.",
          col_widths=[4.2, 2.2])


def build_references(doc):
    heading(doc, "References", 1)
    refs = [
        "Breiman, L. (2001) 'Random forests', Machine Learning, 45(1), pp. 5-32. "
        "doi:10.1023/A:1010933404324.",
        "Chawla, N.V., Bowyer, K.W., Hall, L.O. and Kegelmeyer, W.P. (2002) 'SMOTE: Synthetic "
        "Minority Over-sampling Technique', Journal of Artificial Intelligence Research, 16, "
        "pp. 321-357. doi:10.1613/jair.953.",
        "Chen, T. and Guestrin, C. (2016) 'XGBoost: A scalable tree boosting system', "
        "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and "
        "Data Mining (KDD '16), San Francisco, pp. 785-794. doi:10.1145/2939672.2939785.",
        "Emmanuel, T., Maupong, T., Mpoeleng, D., Semong, T., Mphago, B. and Tabona, O. (2021) "
        "'A survey on missing data in machine learning', Journal of Big Data, 8, Article 140. "
        "doi:10.1186/s40537-021-00516-9.",
        "European Parliament and Council (2016) Regulation (EU) 2016/679 (General Data "
        "Protection Regulation), Article 22. Official Journal of the European Union, L 119.",
        "Garcia, S., Ramirez-Gallego, S., Luengo, J., Benitez, J.M. and Herrera, F. (2016) 'Big "
        "data preprocessing: methods and prospects', Big Data Analytics, 1, Article 9. "
        "doi:10.1186/s41044-016-0014-0.",
        "Goodman, B. and Flaxman, S. (2017) 'European Union regulations on algorithmic "
        "decision-making and a \"right to explanation\"', AI Magazine, 38(3), pp. 50-57. "
        "doi:10.1609/aimag.v38i3.2741.",
        "Hand, D.J. and Henley, W.E. (1997) 'Statistical classification methods in consumer "
        "credit scoring: A review', Journal of the Royal Statistical Society: Series A, 160(3), "
        "pp. 523-541. doi:10.1111/j.1467-985X.1997.00078.x.",
        "He, H. and Garcia, E.A. (2009) 'Learning from imbalanced data', IEEE Transactions on "
        "Knowledge and Data Engineering, 21(9), pp. 1263-1284. doi:10.1109/TKDE.2008.239.",
        "Kozodoi, N., Jacob, J. and Lessmann, S. (2022) 'Fairness in credit scoring: Assessment, "
        "implementation and profit implications', European Journal of Operational Research, "
        "297(3), pp. 1083-1094. doi:10.1016/j.ejor.2021.06.023.",
        "Lessmann, S., Baesens, B., Seow, H.-V. and Thomas, L.C. (2015) 'Benchmarking "
        "state-of-the-art classification algorithms for credit scoring: An update of research', "
        "European Journal of Operational Research, 247(1), pp. 124-136. "
        "doi:10.1016/j.ejor.2015.05.030.",
        "Lundberg, S.M. and Lee, S.-I. (2017) 'A unified approach to interpreting model "
        "predictions', Advances in Neural Information Processing Systems 30 (NIPS 2017), Long "
        "Beach, CA. arXiv:1705.07874.",
        "Sarikaya, A., Correll, M., Bartram, L., Tory, M. and Fisher, D. (2019) 'What do we talk "
        "about when we talk about dashboards?', IEEE Transactions on Visualization and Computer "
        "Graphics, 25(1), pp. 682-692. doi:10.1109/TVCG.2018.2864903.",
        "XGBoost Developers (2022) 'Notes on parameter tuning', XGBoost Documentation. Available "
        "at: https://xgboost.readthedocs.io/en/stable/tutorials/param_tuning.html.",
        "Yang, S., Huang, Z., Xiao, W. and Shen, X. (2025) 'Interpretable credit default "
        "prediction with ensemble learning and SHAP', arXiv preprint arXiv:2505.20815.",
    ]
    for ref in refs:
        p = para(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)


def build_appendices(doc):
    heading(doc, "Appendix A: Additional Exploratory Data Analysis Figures", 1)
    para(doc, "The figures below supplement the exploratory analysis in Chapter 1.")
    figure(doc, os.path.join(EDA_DIR, "eda_amounts.png"),
           "Figure A.1: Distributions of the main financial amount columns, showing strong "
           "right skew.")
    figure(doc, os.path.join(EDA_DIR, "eda_amt_vs_target.png"),
           "Figure A.2: Credit and income amounts compared across repaid and defaulted loans.")
    figure(doc, os.path.join(EDA_DIR, "eda_ext_source.png"),
           "Figure A.3: Distributions of the three external credit scores.")
    figure(doc, os.path.join(EDA_DIR, "eda_age_employment.png"),
           "Figure A.4: Age and years employed against default outcome.")
    figure(doc, os.path.join(EDA_DIR, "eda_missing.png"),
           "Figure A.5: Missing-value footprint by column, dominated by the building survey "
           "block.")
    figure(doc, os.path.join(EDA_DIR, "eda_corr.png"),
           "Figure A.6: Correlation heatmap of the numeric features and TARGET.")

    page_break(doc)
    heading(doc, "Appendix B: Model Evaluation Figures for the Baseline Models", 1)
    para(doc, "The figures below support the hold-out comparison in Section 2.10.")
    figure(doc, os.path.join(FIG_DIR, "cm_logistic_regression.png"),
           "Figure B.1: Confusion matrix of Logistic Regression on the hold-out set.")
    figure(doc, os.path.join(FIG_DIR, "roc_logistic_regression.png"),
           "Figure B.2: ROC curve of Logistic Regression (AUC = 0.749).")
    figure(doc, os.path.join(FIG_DIR, "cm_random_forest.png"),
           "Figure B.3: Confusion matrix of Random Forest on the hold-out set.")
    figure(doc, os.path.join(FIG_DIR, "roc_random_forest.png"),
           "Figure B.4: ROC curve of Random Forest (AUC = 0.742).")

    page_break(doc)
    heading(doc, "Appendix C: System Diagrams", 1)
    para(doc, "The diagrams below document the system architecture described in Chapter 2.")
    figure(doc, os.path.join(SYS_DIR, "dfd_level0.png"),
           "Figure C.1: Level 0 data flow diagram (context diagram) of the prediction system.")
    figure(doc, os.path.join(SYS_DIR, "dfd_level1.png"),
           "Figure C.2: Level 1 data flow diagram showing the pipeline processes and data "
           "stores.")
    figure(doc, os.path.join(SYS_DIR, "erd.png"),
           "Figure C.3: Entity relationship diagram of the Home Credit relational schema.")
    figure(doc, os.path.join(SYS_DIR, "uml_class.png"),
           "Figure C.4: UML class diagram of the pipeline modules.")
    figure(doc, os.path.join(SYS_DIR, "flowchart_preprocessing.png"),
           "Figure C.5: Flowchart of the preprocessing stages.")

    page_break(doc)
    heading(doc, "Appendix D: Project Schedule", 1)
    para(doc, "The Gantt chart below records the project schedule from 13 April 2026 to 28 June "
              "2026, covering planning, data sourcing, the collection plan, EDA, preprocessing, "
              "model development, evaluation, the dashboard, testing and deployment, and final "
              "preparation.")
    figure(doc, os.path.join(WORK_DIR, "gantt_homecredit.png"),
           "Figure D.1: Project Gantt chart (13 April 2026 to 28 June 2026).")


# ----------------------------------------------------------------------
# build + verify
# ----------------------------------------------------------------------

def build():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_page_number_footer(doc)

    build_cover(doc)
    page_break(doc)
    build_toc(doc)
    page_break(doc)
    build_ch1(doc)
    page_break(doc)
    build_ch2(doc)
    page_break(doc)
    build_ch3(doc)
    page_break(doc)
    build_ch4(doc)
    page_break(doc)
    build_ch5(doc)
    page_break(doc)
    build_ch6(doc)
    page_break(doc)
    build_references(doc)
    page_break(doc)
    build_appendices(doc)

    doc.save(OUT_DOCX)
    print("Saved:", OUT_DOCX)


def verify():
    doc = Document(OUT_DOCX)

    # heading tree
    print("\n--- HEADING TREE ---")
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = p.style.name.replace("Heading ", "")
            print(("  " * (int(level) - 1)) + p.text)

    # body word count: Chapter 1 through end of Chapter 5, captions excluded
    print("\n--- WORD COUNT (Chapters 1-5, captions excluded) ---")
    counting = False
    words = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("Chapter 1:"):
            counting = True
        if txt.startswith("Chapter 6:"):
            counting = False
        if counting and txt and p.style.name != "Caption":
            words += len(txt.split())
    print("Main body words:", words)

    # tables
    print("\n--- TABLES ---")
    for i, t in enumerate(doc.tables):
        hdr = " | ".join(c.text for c in t.rows[0].cells)
        print("Table %d (%d rows): %s" % (i + 1, len(t.rows), hdr[:100]))
    headers = ["|".join(c.text for c in t.rows[0].cells) for t in doc.tables]
    assert any("Accuracy" in h and "ROC-AUC" in h for h in headers), "metrics table missing"
    assert any("Training complexity" in h for h in headers), "complexity table missing"
    assert any("Best threshold" in h for h in headers), "threshold table missing"
    assert any("Lead Member" in h for h in headers), "responsibilities table missing"
    print("Required tables: metrics, complexity, threshold, responsibilities -- all present.")

    # images
    n_img = len(doc.inline_shapes)
    print("\nEmbedded images:", n_img, "(expected 24)")
    assert n_img == 24, "image count mismatch"

    # citation check
    full_text = "\n".join(p.text for p in doc.paragraphs)
    cites = ["Emmanuel et al. (2021)", "He and Garcia (2009)", "Yang et al. (2025)",
             "Garcia et al. (2016)", "Hand and Henley, 1997", "Breiman, 2001",
             "Chen and Guestrin, 2016", "XGBoost Developers, 2022", "Chawla et al., 2002",
             "Sarikaya et al. (2019)", "Lessmann et al. (2015)", "Lundberg and Lee, 2017",
             "European Parliament and Council, 2016", "Goodman and Flaxman, 2017",
             "Kozodoi et al., 2022"]
    missing = [c for c in cites if c not in full_text]
    print("In-text citations found:", len(cites) - len(missing), "of", len(cites))
    if missing:
        print("MISSING CITATIONS:", missing)
    assert not missing, "missing in-text citations"
    print("\nVERIFICATION PASSED")


if __name__ == "__main__":
    build()
    verify()
