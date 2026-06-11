# verify_thong_report.py -- post-build checks for ThongWaiKit_5011CEM_CW1_Report.docx
import re
from docx import Document

PATH = r"C:\Users\ASUS\Downloads\07_Team_Handoff\07_Team_Handoff\08_Final_Submission\ThongWaiKit_5011CEM_CW1_Report.docx"

doc = Document(PATH)

print("=== HEADING TREE ===")
for p in doc.paragraphs:
    if p.style.name == "Heading 1":
        print(p.text)
    elif p.style.name == "Heading 2":
        print("   " + p.text)

# Body word count: paragraphs between "Chapter 1" H1 and "References" H1,
# excluding headings and figure captions.
in_body = False
words = 0
em_dash_hits = []
body_text_all = []
for p in doc.paragraphs:
    if p.style.name == "Heading 1":
        if p.text.startswith("Chapter 1"):
            in_body = True
        elif p.text.strip() == "References":
            in_body = False
    if in_body:
        body_text_all.append(p.text)
        if p.style.name.startswith("Heading"):
            continue
        if p.text.strip().startswith("Figure "):
            continue
        words += len(p.text.split())

print()
print("=== BODY WORD COUNT (Ch1-Ch4, excl. headings/captions):", words, "===")

# Em-dash / non-ASCII scan over the whole document
full_text = "\n".join(p.text for p in doc.paragraphs)
assert "—" not in full_text, "EM DASH FOUND"
assert "–" not in full_text, "EN DASH FOUND"
non_ascii = sorted(set(ch for ch in full_text if ord(ch) > 127))
print("Em-dash check: PASS. Non-ASCII chars in doc:", non_ascii if non_ascii else "none")

# Embedded images count
img_count = len(doc.inline_shapes)
print("Embedded inline images:", img_count, "(expected 6)")
assert img_count == 6, "wrong image count"

# Citation count (unique sources cited in Ch1-Ch4 body text)
body_str = " ".join(body_text_all)
cites = {
    "Lessmann et al.": r"Lessmann et al\.,? \(?2015|Lessmann et al\., 2015",
    "Breiman": r"Breiman,? \(?2001",
    "Chen and Guestrin": r"Chen and Guestrin",
    "Xia et al.": r"Xia et al\.",
    "He and Garcia": r"He and Garcia",
    "Chawla et al.": r"Chawla et al\.",
    "XGBoost Developers": r"XGBoost Developers, 2022",
    "Garcia et al. 2016": r"Garcia et al\., 2016",
    "Al-qerem et al.": r"Al-qerem",
    "Yang et al.": r"Yang et al\.",
    "Li and Wu": r"Li and Wu",
    "Lundberg and Lee": r"Lundberg and Lee",
    "GDPR Art22": r"European Parliament and Council, 2016",
    "Kozodoi": r"Kozodoi",
}
total_instances = 0
cited = []
for name, pat in cites.items():
    n = len(re.findall(pat, body_str))
    if n:
        cited.append((name, n))
        total_instances += n
print()
print("=== IN-TEXT CITATIONS (Ch1-Ch4) ===")
for name, n in cited:
    print(f"  {name}: {n}")
print("Unique sources cited:", len(cited), "| total in-text instances:", total_instances)
assert len(cited) >= 12, "fewer than 12 unique citations"

# TOC field present?
xml = doc.element.xml
assert "TOC \\o" in xml, "TOC field missing"
print()
print("TOC field: PRESENT")
print("References paragraphs:", sum(1 for p in doc.paragraphs
      if p.paragraph_format.first_line_indent is not None
      and p.paragraph_format.first_line_indent < 0))
print("ALL CHECKS PASSED")
