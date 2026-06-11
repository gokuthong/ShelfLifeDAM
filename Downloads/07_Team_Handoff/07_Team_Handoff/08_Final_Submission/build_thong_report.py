# build_thong_report.py
# Generates ThongWaiKit_5011CEM_CW1_Report.docx (Individual Progressive Report).
# Run with:  python -X utf8 build_thong_report.py
# Source of truth: 06_Working_Files/REPORT_FACTS_PACK.md and verified_references.md

import os
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = r"C:\Users\ASUS\Downloads\07_Team_Handoff\07_Team_Handoff"
OUT_DIR = os.path.join(BASE, "08_Final_Submission")
OUT_DOCX = os.path.join(OUT_DIR, "ThongWaiKit_5011CEM_CW1_Report.docx")

IMG = {
    "flowchart": os.path.join(BASE, r"04_Diagrams\System_Diagrams\flowchart_preprocessing.png"),
    "uml": os.path.join(BASE, r"04_Diagrams\System_Diagrams\uml_class.png"),
    "erd": os.path.join(BASE, r"04_Diagrams\System_Diagrams\erd.png"),
    "eda_missing": os.path.join(BASE, r"04_Diagrams\EDA_Plots\eda_missing.png"),
    "eda_ext_source": os.path.join(BASE, r"04_Diagrams\EDA_Plots\eda_ext_source.png"),
    "gantt": os.path.join(BASE, r"06_Working_Files\gantt_homecredit.png"),
}

for key, path in IMG.items():
    assert os.path.exists(path), "Missing image: " + path

# ----------------------------------------------------------------------------
# Report content. Item types:
#   h1 = chapter heading (Heading 1), h2 = numbered subsection (Heading 2),
#   p  = body paragraph, li = bullet item, fig = (imgkey, caption, width_in)
#   pagebreak = explicit page break before next item
# ----------------------------------------------------------------------------

CH1 = [
    ("h1", "Chapter 1: Introduction"),
    ("h2", "1.1 Project Background and Motivation"),
    ("p",
     "Consumer lenders process very large volumes of loan applications, and every approval "
     "decision carries the risk that the borrower fails to repay. My team and I selected the "
     "Home Credit Default Risk dataset, released by the Home Credit Group for a public Kaggle "
     "competition in 2018, as the foundation for our 5011CEM project on loan-default prediction. "
     "The main table records 307,511 historical applications described by 122 attributes, which "
     "comfortably satisfies the module requirement of at least 100,000 records and places the "
     "work firmly in big-data territory. From my position as the member responsible for data "
     "engineering and modelling, the dataset is attractive for two technical reasons. Firstly, "
     "only 8.07 percent of the applications ended in payment difficulty, so any model must learn "
     "from a heavily imbalanced label distribution. Secondly, roughly 24 percent of all cells in "
     "the table are missing, which turns preprocessing quality into a central design problem "
     "rather than a routine clean-up step. These two properties shape almost every decision "
     "described in this report."),
    ("h2", "1.2 Problem Statement"),
    ("p",
     "The core problem is to predict, at application time, whether an applicant will experience "
     "payment difficulty using only the attributes captured on the application record. A trivial "
     "classifier that approves everybody would score about 91.9 percent accuracy on this data "
     "while detecting zero defaulters, so plain accuracy is a misleading objective. The real "
     "challenge, and the focus of my individual contribution, is to design a preprocessing and "
     "modelling pipeline that extracts a dependable ranking signal from incomplete, skewed and "
     "imbalanced data, and to do so with algorithms that remain computationally efficient on a "
     "table of more than three hundred thousand rows."),
    ("h2", "1.3 Aim and Objectives"),
    ("p",
     "The aim of my portion of the project is to design, and progressively implement, an "
     "efficient and leakage-safe machine learning pipeline that predicts loan default on the "
     "Home Credit data. In order to achieve this aim, I set five objectives for myself."),
    ("li",
     "To audit the data quality problems in the raw table, including missing values, sentinel "
     "codes and skewed distributions, and to design a preprocessing plan that repairs them "
     "without leaking information from the test partition."),
    ("li",
     "To gather statistical evidence, through chi-square tests and point-biserial correlations, "
     "that justifies which variables deserve a place in the models."),
    ("li",
     "To plan a controlled comparison between a linear baseline and gradient-boosted trees, "
     "with XGBoost as the primary candidate."),
    ("li",
     "To design a dedicated experiment that selects one, and only one, class imbalance "
     "correction strategy."),
    ("li",
     "To define an evaluation protocol built on imbalance-aware metrics instead of raw "
     "accuracy."),
    ("h2", "1.4 Scope and Division of Work"),
    ("p",
     "The project scope covers the single application_train table; the six relational child "
     "tables that join on the application identifier are documented in the design but reserved "
     "as future work, because a well-controlled single-table pipeline must come first. Within "
     "the two-person team, my teammate Bryan Tey leads the exploratory data analysis, the "
     "Logistic Regression and Random Forest models, the dashboard Overview and Live Prediction "
     "pages, and the assembly of the group report. I lead the data loading and preprocessing "
     "modules, the XGBoost model together with the imbalance experiment, the shared evaluation "
     "module, the dashboard Model Comparison page, and the architecture documentation. Joint "
     "decisions, such as the dataset choice and the evaluation metrics, were made by my team "
     "and I together."),
    ("h2", "1.5 Success Criteria"),
    ("p",
     "I will treat my work as successful if four conditions hold at the end of the project. "
     "The pipeline must be fully reproducible from one command with a fixed random seed. The "
     "boosted model must beat the linear baseline on F1 and ROC-AUC measured on a stratified "
     "hold-out partition. The chosen imbalance correction must be backed by a controlled "
     "experiment rather than habit. Lastly, every variable that enters the models must be "
     "supported by a statistical test or a documented engineering reason."),
    ("h2", "1.6 Report Organisation"),
    ("p",
     "This progressive report is organised as follows. Chapter 2 reviews the literature on "
     "credit-risk modelling, gradient boosting, class imbalance and large-scale preprocessing, "
     "and derives my research questions and hypotheses. Chapter 3 presents the data source, the "
     "quality issues I have already confirmed, the preprocessing pipeline design and the storage "
     "plan. Chapter 4 describes the statistical evidence collected so far, the candidate "
     "algorithms, and the implementation and evaluation plan for the remaining weeks. The "
     "appendices contain the project Gantt chart and the main design diagrams."),
]

CH2 = [
    ("h1", "Chapter 2: Literature Review"),
    ("p",
     "This chapter reviews the published work that informs my technical design. I begin with "
     "the movement from linear scorecards to tree ensembles, then narrow down to gradient "
     "boosting, class imbalance and preprocessing at scale, before closing with studies on the "
     "same Home Credit dataset and the regulatory constraints that follow any automated credit "
     "decision. The research questions and hypotheses at the end of the chapter are derived "
     "directly from this review."),
    ("h2", "2.1 From Linear Scorecards to Tree Ensembles"),
    ("p",
     "Credit scoring has a long statistical tradition, and for decades the industry default was "
     "a linear or logistic scorecard. The most thorough modern benchmark remains the study by "
     "Lessmann et al. (2015), which compared 41 classification algorithms across multiple real "
     "credit datasets and found that ensemble methods, including random forests, consistently "
     "outperformed logistic regression on discrimination measures. This result matters for my "
     "design in two ways. Firstly, it justifies keeping a logistic model in the project purely "
     "as a baseline, since any ensemble that cannot beat it is not earning its complexity. "
     "Secondly, it warns that the margin of improvement is often moderate, so the comparison "
     "must use sensitive metrics. The random forest itself, introduced by Breiman (2001), "
     "remains a strong reference ensemble because bootstrap aggregation with random feature "
     "selection gives robustness to overfitting with very little tuning."),
    ("p",
     "Reading this body of work as a data engineer rather than as a statistician, I take away "
     "a procedural rule: the comparison protocol matters as much as the contenders. A "
     "benchmark is only trustworthy when every algorithm consumes the same preprocessed "
     "matrix, the same partitions and the same metrics. Several of the differences reported "
     "between studies dissolve once the experimental setup is held constant, and that "
     "observation is the reason my design binds all models to one shared preprocessing "
     "transformer instead of letting each model carry its own."),
    ("h2", "2.2 Gradient Boosting for Credit Risk"),
    ("p",
     "My primary modelling interest is gradient boosting. Chen and Guestrin (2016) presented "
     "XGBoost as a regularised boosting system with sparsity-aware split finding and "
     "out-of-core computation, and demonstrated that it scales to billions of examples. Two of "
     "its engineering properties map directly onto the Home Credit data. The sparsity-aware "
     "algorithm treats missing values as a first-class case by learning a default split "
     "direction, which suits a table where around a quarter of the cells are empty, and the "
     "histogram-based tree method bins continuous features once, which keeps training time low "
     "on wide matrices. In the credit domain specifically, Xia et al. (2017) showed that a "
     "boosted tree approach with careful hyper-parameter optimisation outperformed standard "
     "benchmarks on several credit scoring datasets. Their finding that tuning matters as much "
     "as the algorithm choice motivates my plan to fix a small set of sensible parameters early "
     "and to spend the remaining effort on threshold and imbalance decisions instead of an "
     "endless grid search."),
    ("h2", "2.3 Learning from Imbalanced Data"),
    ("p",
     "With only 8.07 percent positive labels, the Home Credit task sits squarely in the "
     "imbalanced learning literature. He and Garcia (2009) surveyed the field and grouped the "
     "remedies into sampling methods, cost-sensitive learning and decision-threshold "
     "adjustment, while stressing that accuracy is the wrong yardstick whenever one class "
     "dominates. The best-known sampling method is SMOTE by Chawla et al. (2002), which "
     "synthesises new minority examples by interpolating between neighbours rather than simply "
     "duplicating rows. On the cost-sensitive side, the official XGBoost guidance recommends "
     "the scale_pos_weight parameter, which reweights the gradient of the minority class, "
     "together with ranking-based evaluation metrics such as AUC (XGBoost Developers, 2022). "
     "What the literature does not settle is which single correction works best for a given "
     "dataset, and several authors caution that stacking corrections can over-shoot, pushing "
     "recall up while precision collapses. This open question becomes my second research "
     "question below, and I have designed a dedicated experiment around it."),
    ("p",
     "A further point from He and Garcia (2009) deserves emphasis because it changes how "
     "results must be read. Under imbalance, the precision of the minority class is bounded "
     "by the base rate: even a model with respectable recall will produce many false alarms "
     "when only one applicant in twelve actually defaults. Consequently, a credit model "
     "should be judged on the trade-off curve it offers, and on where a chosen threshold "
     "places the business along that curve, not on any single headline number. This insight "
     "directly shapes the evaluation plan in Chapter 4, which treats threshold selection as "
     "a first-class optimisation step rather than an afterthought."),
    ("h2", "2.4 Preprocessing and Data Quality at Scale"),
    ("p",
     "Garcia et al. (2016) argue that preprocessing, covering cleaning, imputation, feature "
     "selection and imbalance treatment, is the dominant cost in big-data analytics and that "
     "poorly planned preprocessing silently destroys model quality. Their framing matches my "
     "experience with this dataset, where the preprocessing code is larger than the modelling "
     "code. Closer to my exact problem, Al-qerem, Al-Naymat and Alhasan (2019) worked on the "
     "Home Credit data itself and reported that comprehensive preprocessing combined with "
     "feature selection improved classifier performance by a large margin, in some "
     "configurations up to around forty percent. Their result is direct evidence that on this "
     "particular dataset the preprocessing decisions are not cosmetic, and it underpins my "
     "third research question about which decisions matter most. Moreover, both papers agree "
     "on a discipline point that is easy to violate in practice: every statistic used to "
     "repair the data, whether a median, a mode or a scaling factor, must be learned from the "
     "training partition alone. Fitting an imputer on the full table quietly transfers "
     "information from the test rows into the training process, and the resulting evaluation "
     "flatters the model. My pipeline design in Chapter 3 enforces this rule structurally, by "
     "placing every repair inside a transformer that is fitted only on training data."),
    ("h2", "2.5 Prior Work on the Home Credit Dataset"),
    ("p",
     "Two recent studies frame what a realistic outcome looks like. Yang et al. (2025) trained "
     "logistic regression, random forest and boosted ensembles on the Home Credit data and "
     "found that the ensembles won, with the external credit bureau scores dominating the SHAP "
     "attributions. Li and Wu (2024) compared nine models for loan default and similarly "
     "concluded that tree ensembles were the most efficient and stable, with explainable "
     "attributions identifying a compact set of key drivers. Both studies reassure me that the "
     "team's three-model line-up is sensible, and they set an expectation that the external "
     "source scores will top my own importance rankings, which is exactly what my early "
     "correlation tests in Chapter 4 suggest."),
    ("h2", "2.6 Explainability and Regulatory Constraints"),
    ("p",
     "A credit model that affects real applicants cannot be a pure black box. Lundberg and Lee "
     "(2017) unified several attribution methods under the SHAP framework, which the project "
     "plans to reference when explaining individual predictions. On the regulatory side, "
     "Article 22 of the GDPR grants individuals the right not to be subject to solely automated "
     "decisions with significant effects (European Parliament and Council, 2016), and Kozodoi, "
     "Jacob and Lessmann (2022) show that fairness constraints in credit scoring interact with "
     "profit in measurable ways. For my design the practical consequence is that demographic "
     "attributes such as gender, although statistically predictive in this dataset, must be "
     "flagged as an ethical limitation rather than celebrated as signal."),
    ("h2", "2.7 Research Questions and Hypotheses"),
    ("p",
     "The review above leaves three questions that my work will answer on this dataset. "
     "RQ1 asks whether gradient boosting genuinely outperforms a linear baseline on imbalanced "
     "credit data once both receive a fair imbalance correction. RQ2 asks which single "
     "imbalance correction, cost weighting through scale_pos_weight or SMOTE oversampling, is "
     "more effective for XGBoost on this data. RQ3 asks which preprocessing decisions matter "
     "most for a real dataset where roughly 24 percent of cells are missing."),
    ("p",
     "Each question carries a testable hypothesis. H1 states that XGBoost will achieve a higher "
     "F1 score and ROC-AUC than logistic regression on the stratified hold-out, in line with "
     "the benchmark evidence (Lessmann et al., 2015; Yang et al., 2025). H2 states that cost "
     "weighting alone will yield a better F1 than SMOTE oversampling at the default 0.5 "
     "threshold, and that combining the two will over-correct, raising recall while crushing "
     "precision; this follows the official tuning guidance (XGBoost Developers, 2022) and will "
     "be tested by the four-arm experiment described in Chapter 4. H3 states that the largest "
     "preprocessing gains will come from leakage-safe imputation and from removing "
     "near-empty columns, consistent with the findings of Al-qerem, Al-Naymat and Alhasan "
     "(2019). The remaining chapters describe the design that will put these hypotheses to "
     "the test."),
]

CH3 = [
    ("h1", "Chapter 3: Design and Planning of Project and Data Analysis"),
    ("p",
     "This chapter records the part of the project that is already concrete: the data source "
     "and its schema, the quality problems I have verified in the raw file, the preprocessing "
     "pipeline I have designed and largely implemented, the system architecture, and the "
     "storage plan."),
    ("h2", "3.1 Data Source, Schema and Format"),
    ("p",
     "The data source is the Home Credit Default Risk dataset from the public Kaggle "
     "competition platform. The core file, application_train.csv, is a flat comma-separated "
     "table of 307,511 rows and 122 columns occupying about 166 MB on disk, with one row per "
     "loan application identified by SK_ID_CURR. The label column TARGET equals 1 when the "
     "applicant had payment difficulty and 0 otherwise, giving 24,825 positive cases, a default "
     "rate of 8.07 percent. The columns mix numeric attributes, such as credit amount, annuity "
     "and the three external bureau scores, with categorical attributes such as income type, "
     "education and contract type; a supplied data dictionary file documents every field. Six "
     "child tables, including bureau records and previous applications, join to the core table "
     "through SK_ID_CURR in a classic one-to-many relational layout, and the entity "
     "relationship diagram in Appendix B captures this structure. The deployed pipeline uses "
     "only the core table, while the children are mapped in the design as the documented "
     "extension path. The data is released by the lender in anonymised form for public "
     "research, contains no direct identifiers, and exceeds the assignment requirement of one "
     "hundred thousand records."),
    ("h2", "3.2 Data Quality Issues Confirmed So Far"),
    ("p",
     "I have completed a quality audit of the raw table, and three issues stand out. The first "
     "is missingness on a scale that rules out row deletion: roughly 24 percent of all cells "
     "are empty, and the damage is concentrated in a block of building-survey columns, such as "
     "the common-area and non-living-apartment measurements, that are mostly blank. Figure 3.1 "
     "shows the missing-value footprint across the worst columns. Dropping every incomplete "
     "row would discard the majority of the dataset, so imputation is unavoidable."),
    ("fig", ("eda_missing",
             "Figure 3.1: Missing-value footprint of the worst-affected columns in "
             "application_train.", 6.0)),
    ("p",
     "The second issue is a sentinel code. The DAYS_EMPLOYED column uses the value 365243, "
     "which translates to roughly one thousand years of employment, as a placeholder for "
     "pensioners and unemployed applicants. The code appears in 55,374 rows, which is 18.0 "
     "percent of the table, and any statistic computed over the raw column is meaningless "
     "until the sentinel is replaced with a proper missing marker. The third issue is heavy "
     "right skew in the monetary columns; incomes and credit amounts have long upper tails, "
     "which my teammate's exploratory plots confirm on a log scale. Besides that, the audit "
     "found zero exact duplicate rows, which is one less repair to perform."),
    ("h2", "3.3 Preprocessing Pipeline Design"),
    ("p",
     "Figure 3.2 presents the preprocessing flowchart I designed for the project; an enlarged "
     "copy appears in Appendix B. The pipeline runs in five stages. Stage one loads the CSV, "
     "immediately converts the DAYS_EMPLOYED sentinel to a missing value, and caches the frame "
     "in parquet format so later runs reload several times faster. Stage two cleans the data "
     "by dropping exact duplicates and logging the missing footprint, while deliberately "
     "keeping incomplete rows for later imputation. Stage three engineers six domain ratios, "
     "including the credit-to-income ratio, the annuity-to-income ratio, the credit term and "
     "an age variable derived from the birth-days column, with any division by zero mapped to "
     "a missing value. Stage four selects feature columns by removing the identifier and the "
     "label, removing every column that is at least 60 percent missing, and removing the "
     "twenty low-information document flags. Stage five is a column transformer that imputes "
     "numeric columns with the median and scales them, and imputes categorical columns with "
     "the mode before one-hot encoding with unknown categories ignored."),
    ("fig", ("flowchart", "Figure 3.2: Preprocessing pipeline flowchart for the Home Credit "
             "data.", 5.4)),
    ("p",
     "Two design rules run through every stage. The first is leakage safety: the column "
     "transformer is fitted on the training partition only, so the medians, modes and scaling "
     "parameters never see test rows, a discipline the preprocessing literature treats as "
     "essential at scale (Garcia et al., 2016). The second is reproducibility: the split and "
     "every model seed are fixed at 42, so any team member can regenerate identical partitions. "
     "The stratified 80/20 split preserves the 8.07 percent default rate on both sides, "
     "producing 246,008 training rows and 61,503 test rows that remain untouched until final "
     "evaluation."),
    ("h2", "3.4 System Architecture and Functions"),
    ("p",
     "The system is organised as a small set of single-purpose Python modules, and the UML "
     "class diagram in Appendix B documents their relationships. A data loader handles the "
     "CSV-to-parquet caching and the sentinel repair. The preprocessing module owns the "
     "cleaning, feature engineering and the column transformer. Each model lives behind a "
     "common interface inside one scikit-learn pipeline object, so the preprocessor and the "
     "estimator travel together and cannot drift apart. A shared evaluation module, which I "
     "maintain, computes every metric and writes a single consolidated metrics file; the "
     "Streamlit dashboard reads that same file, which guarantees the report numbers and the "
     "dashboard numbers can never disagree. A single orchestration script runs the whole chain "
     "from raw CSV to figures, and a pytest suite covers the loading, preprocessing and "
     "evaluation functions. In a nutshell, the architecture is deliberately boring, because "
     "boring architectures are reproducible."),
    ("h2", "3.5 Data Storage Considerations"),
    ("p",
     "Storage looks trivial for one table but stops being trivial at the project's full "
     "extent. The 166 MB CSV loads slowly in pandas, so the parquet cache is the pragmatic "
     "fix for the current single-table scope, cutting reload time several-fold while keeping "
     "the data in ordinary files. The complete relational set, however, spans seven tables and "
     "roughly 2.7 GB, and repeatedly joining it in memory would be wasteful. For that "
     "extension my team and I have discussed a document store such as MongoDB, which we used "
     "in the module labs, as a way to hold semi-structured application records with their "
     "child histories embedded per applicant. For the progressive scope, flat CSV plus parquet "
     "remains adequate for single-table batch analysis, and the trade-off is recorded here so "
     "the group report can revisit it with measurements."),
]

CH4 = [
    ("h1", "Chapter 4: Model Design and Planning"),
    ("p",
     "This chapter moves from data to models. It records the analysis tools, the statistical "
     "evidence I have already computed to justify the variables, the algorithms the project "
     "will train, the imbalance handling design, and the implementation and evaluation plan "
     "for the remaining weeks. Following the progressive nature of this report, the statistics "
     "are reported as completed work, while the modelling is described in planning voice."),
    ("h2", "4.1 Analysis Tools and Techniques"),
    ("p",
     "The project standardises on Python 3.12 with pandas for data handling, scikit-learn for "
     "the preprocessing transformer, the baseline models and the metrics, the XGBoost library "
     "for gradient boosting, imbalanced-learn for the SMOTE arm of the experiment, and "
     "matplotlib with seaborn for static figures. Jupyter Notebook is used for exploration and "
     "VS Code for the production modules, with the dashboard built in Streamlit. I chose this "
     "stack because every component is open source, widely documented and reproducible on both "
     "team members' machines, and because the scikit-learn pipeline object lets me bind the "
     "preprocessor to each estimator so that no preprocessing step can be applied "
     "inconsistently between models."),
    ("h2", "4.2 Statistical Evidence for Variable Selection"),
    ("p",
     "Before any model is trained, every candidate variable should justify itself "
     "statistically. For the categorical attributes I ran chi-square tests of independence "
     "against the default label. The strongest association belongs to income type with a "
     "chi-square statistic of 1253.5 on 7 degrees of freedom, followed by education type at "
     "1019.2 on 4 degrees, gender at 920.8 on 2 degrees, family status at 504.7 on 5 degrees "
     "and contract type at 293.2 on 1 degree, with every test significant at p below 0.001. "
     "These results confirm that the application-form categories carry genuine signal and "
     "deserve their one-hot encoded columns."),
    ("p",
     "For the numeric attributes I computed point-biserial correlations with the label. The "
     "three external bureau scores dominate: EXT_SOURCE_3 correlates at -0.179, EXT_SOURCE_2 "
     "at -0.161 and EXT_SOURCE_1 at -0.155, all significant at p below 0.05, meaning lower "
     "external scores accompany higher default risk. Figure 4.1 shows their distributions. "
     "Age and years employed follow at -0.078 and -0.075, so younger and shorter-employed "
     "applicants default more often. Notably, raw income correlates at only -0.004, which is "
     "practically nothing, while the engineered annuity-to-income ratio and credit term carry "
     "small positive correlations. The lesson I draw is that affordability ratios, not raw "
     "income, encode repayment stress, which retrospectively validates the feature "
     "engineering stage of the pipeline. The dominance of the external scores also matches "
     "the SHAP findings reported by Yang et al. (2025) on the same data."),
    ("fig", ("eda_ext_source",
             "Figure 4.1: Distributions of the three external bureau scores, the strongest "
             "single predictors of default.", 6.0)),
    ("h2", "4.3 Candidate Algorithms and Justification"),
    ("p",
     "The project will train three classifiers inside the shared pipeline, all consuming the "
     "identical preprocessed matrix. The first is Logistic Regression with the liblinear "
     "solver, balanced class weights and up to 500 iterations. It exists as the interpretable "
     "industry baseline, and under H1 its job is to set the bar that the ensembles must clear. "
     "The second is a Random Forest of 200 trees with maximum depth 20, a minimum of 10 "
     "samples per leaf and balanced subsample weighting; bagging with random feature selection "
     "provides a robust mid-point between the linear model and boosting (Breiman, 2001). The "
     "third, and my primary candidate, is XGBoost configured with 400 trees of maximum depth "
     "6, a learning rate of 0.1, row and column subsampling at 0.9, the histogram tree method "
     "and the area under the precision-recall curve as the evaluation metric."),
    ("p",
     "My justification for putting gradient boosting first is partly empirical and partly "
     "algorithmic. Empirically, the benchmark literature and the Home Credit studies "
     "consistently rank boosted ensembles at the top for credit discrimination (Lessmann et "
     "al., 2015; Li and Wu, 2024). Algorithmically, the sparsity-aware split finding of "
     "XGBoost handles the dataset's pervasive missingness natively, and the histogram method "
     "bins features once before training, so the expected training complexity grows linearly "
     "with the number of rows for a fixed tree budget (Chen and Guestrin, 2016). I therefore "
     "expect XGBoost to train faster than the single-threaded liblinear solver on the wide "
     "one-hot matrix despite carrying the most trees, and the project will record measured "
     "fit and predict times for all three models to verify this expectation rather than "
     "assert it."),
    ("h2", "4.4 Imbalance Handling Design"),
    ("p",
     "RQ2 receives its own experiment, implemented as a four-arm comparison on the same "
     "stratified hold-out. Arm one is an uncorrected XGBoost with scale_pos_weight equal to 1, "
     "which establishes how badly the raw imbalance hurts. Arm two applies cost weighting "
     "only, setting scale_pos_weight to the ratio of negative to positive training cases, "
     "which is approximately 11.4 for this data, following the official guidance (XGBoost "
     "Developers, 2022). Arms three and four replace weighting with SMOTE oversampling "
     "(Chawla et al., 2002) at sampling ratios of 0.3 and 0.5 respectively. Under H2, I "
     "predict that the weighting-only arm will deliver the best F1 at the default threshold, "
     "and that any attempt to combine SMOTE with weighting would double-correct the class "
     "balance, inflating recall while precision collapses, which is why the design keeps the "
     "corrections strictly separate. Whatever the outcome, the deployed model will carry "
     "exactly one correction, selected by this experiment, and the losing arms will be "
     "reported honestly in the final group report."),
    ("h2", "4.5 Implementation and Evaluation Plan"),
    ("p",
     "Evaluation is designed around the warning of He and Garcia (2009) that accuracy "
     "misleads under imbalance. The primary metrics will be precision, recall, F1 and ROC-AUC "
     "on the untouched 61,503-row hold-out, with the confusion matrix reported in full for "
     "each model so the cost of every error type stays visible. Stratified K-fold "
     "cross-validation on the training partition will check that the hold-out numbers are not "
     "a lucky draw, and the fixed seed keeps each fold reproducible. After the main "
     "comparison, the plan includes a decision-threshold sweep for every model, because the "
     "default 0.5 cut-off is rarely optimal for an 8 percent positive class and threshold "
     "tuning is the cheapest optimisation available. The plan also schedules calibration "
     "checks, since imbalance corrections are known to distort raw probability outputs, and "
     "if the scores prove poorly calibrated the dashboard will present percentile risk bands "
     "instead of quoting raw probabilities."),
    ("p",
     "Beyond the metrics, the implementation plan assigns explainability and ethics work. "
     "Feature attributions will be produced for the final model, with SHAP as the planned "
     "framework (Lundberg and Lee, 2017), so that individual predictions can be explained to "
     "a reviewer. This matters legally as well as technically, because Article 22 of the GDPR "
     "limits solely automated credit decisions (European Parliament and Council, 2016), and "
     "the fairness literature shows that predictive demographic attributes demand explicit "
     "scrutiny before deployment (Kozodoi, Jacob and Lessmann, 2022). Gender is a concrete "
     "example in this dataset: it tests as predictive, yet relying on it would be ethically "
     "and legally questionable, so it is flagged as a limitation for the final report to "
     "address. Lastly, the engineering plan continues the practices already in place, namely "
     "trunk-based version control on the team GitHub repository with descriptive commit "
     "messages split evenly between members, pytest unit tests for the data and evaluation "
     "modules, and the Gantt schedule in Appendix A, which carries the project from April "
     "13 to June 28, 2026. In a nutshell, the design fixes the data foundation first, lets "
     "statistics nominate the variables, and reserves the model competition for a fair, "
     "single-corrected, threshold-tuned comparison."),
]

REFERENCES = [
    "Al-qerem, A., Al-Naymat, G. and Alhasan, M. (2019) 'Loan default prediction model "
    "improvement through comprehensive preprocessing and features selection', 2019 "
    "International Arab Conference on Information Technology (ACIT), Al Ain, UAE, pp. 235-240. "
    "doi:10.1109/ACIT47987.2019.8991084.",

    "Breiman, L. (2001) 'Random forests', Machine Learning, 45(1), pp. 5-32. "
    "doi:10.1023/A:1010933404324.",

    "Chawla, N.V., Bowyer, K.W., Hall, L.O. and Kegelmeyer, W.P. (2002) 'SMOTE: Synthetic "
    "Minority Over-sampling Technique', Journal of Artificial Intelligence Research, 16, "
    "pp. 321-357. doi:10.1613/jair.953.",

    "Chen, T. and Guestrin, C. (2016) 'XGBoost: A scalable tree boosting system', Proceedings "
    "of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining "
    "(KDD '16), San Francisco, pp. 785-794. doi:10.1145/2939672.2939785.",

    "European Parliament and Council (2016) Regulation (EU) 2016/679 (General Data Protection "
    "Regulation), Article 22. Official Journal of the European Union, L 119.",

    "Garcia, S., Ramirez-Gallego, S., Luengo, J., Benitez, J.M. and Herrera, F. (2016) 'Big "
    "data preprocessing: methods and prospects', Big Data Analytics, 1, Article 9. "
    "doi:10.1186/s41044-016-0014-0.",

    "He, H. and Garcia, E.A. (2009) 'Learning from imbalanced data', IEEE Transactions on "
    "Knowledge and Data Engineering, 21(9), pp. 1263-1284. doi:10.1109/TKDE.2008.239.",

    "Kozodoi, N., Jacob, J. and Lessmann, S. (2022) 'Fairness in credit scoring: Assessment, "
    "implementation and profit implications', European Journal of Operational Research, "
    "297(3), pp. 1083-1094. doi:10.1016/j.ejor.2021.06.023.",

    "Lessmann, S., Baesens, B., Seow, H.-V. and Thomas, L.C. (2015) 'Benchmarking "
    "state-of-the-art classification algorithms for credit scoring: An update of research', "
    "European Journal of Operational Research, 247(1), pp. 124-136. "
    "doi:10.1016/j.ejor.2015.05.030.",

    "Li, H. and Wu, W. (2024) 'Loan default predictability with explainable machine "
    "learning', Finance Research Letters, 60, 104867. doi:10.1016/j.frl.2023.104867.",

    "Lundberg, S.M. and Lee, S.-I. (2017) 'A unified approach to interpreting model "
    "predictions', Advances in Neural Information Processing Systems 30 (NIPS 2017), Long "
    "Beach, CA. arXiv:1705.07874.",

    "Xia, Y., Liu, C., Li, Y. and Liu, N. (2017) 'A boosted decision tree approach using "
    "Bayesian hyper-parameter optimization for credit scoring', Expert Systems with "
    "Applications, 78, pp. 225-241. doi:10.1016/j.eswa.2017.02.017.",

    "XGBoost Developers (2022) 'Notes on parameter tuning', XGBoost Documentation. Available "
    "at: https://xgboost.readthedocs.io/en/stable/tutorials/param_tuning.html.",

    "Yang, S., Huang, Z., Xiao, W. and Shen, X. (2025) 'Interpretable credit default "
    "prediction with ensemble learning and SHAP', arXiv preprint arXiv:2505.20815.",
]

APPENDICES = [
    ("h1", "Appendices"),
    ("h2", "Appendix A: Project Gantt Chart"),
    ("p",
     "Figure A.1 shows the project schedule from 13 April 2026 to 28 June 2026, covering "
     "planning, data sourcing, the collection plan, exploratory analysis, preprocessing, "
     "model development, evaluation, the dashboard, testing and deployment, and final report "
     "preparation."),
    ("fig", ("gantt", "Figure A.1: Project Gantt chart, 13 April 2026 to 28 June 2026.", 6.0)),
    ("h2", "Appendix B: System Design Diagrams"),
    ("p",
     "Figure B.1 reproduces the preprocessing flowchart from Chapter 3 at full size. Figure "
     "B.2 presents the UML class diagram of the pipeline modules described in Section 3.4, "
     "and Figure B.3 presents the entity relationship diagram of the full Home Credit "
     "relational schema referenced in Section 3.1."),
    ("fig", ("flowchart",
             "Figure B.1: Preprocessing pipeline flowchart (enlarged copy of Figure 3.2).",
             6.0)),
    ("fig", ("uml", "Figure B.2: UML class diagram of the pipeline modules.", 6.0)),
    ("fig", ("erd", "Figure B.3: Entity relationship diagram of the Home Credit relational "
             "schema.", 6.0)),
]

# ----------------------------------------------------------------------------
# docx helpers
# ----------------------------------------------------------------------------

def set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    specs = {"Heading 1": (16, True), "Heading 2": (13, True), "Heading 3": (12, True)}
    for name, (size, bold) in specs.items():
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.5

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Times New Roman"
    bullet.font.size = Pt(12)
    bullet.paragraph_format.line_spacing = 1.5


def add_page_break(doc):
    doc.add_page_break()


def add_cover_sheet(doc):
    def centered(text, size=12, bold=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p

    centered("INTI International College Penang", 14, True)
    centered("School of Computing", 14, True)
    centered("3+0 Bachelor of Science (Hons) in Computer Science, "
             "in collaboration with Coventry University, UK", 12, False)
    centered("Coursework cover sheet", 13, True, 12)
    p = doc.add_paragraph()
    run = p.add_run("Section A - To be completed by the student")
    run.bold = True

    rows = [
        ("Full Name", "Thong Wai Kit"),
        ("CU Student ID Number", "P23015668"),
        ("Semester", "April 2026"),
        ("Lecturer", "Ms. Vimala Doraisamy"),
        ("Module Code and Title", "5011CEM Big Data Programming Project"),
        ("Assignment No. / Title",
         "Home Credit Loan-Default Prediction - 100% of Module Mark"),
        ("Hand out date", "Week 1"),
        ("Due date",
         "Individual Progress Report - WK10 / Final Report - WK12 / VIVA - WK13 & 14"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        run = c0.paragraphs[0].add_run(label)
        run.bold = True
        c1.text = value
        c0.width = Inches(2.3)
        c1.width = Inches(4.0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Penalties: ")
    run.bold = True
    p.add_run(
        "No late work will be accepted. If you are unable to submit coursework on time due "
        "to extenuating circumstances, you may be eligible for an extension. Please consult "
        "the lecturer.")

    p = doc.add_paragraph()
    run = p.add_run("Declaration: ")
    run.bold = True
    p.add_run(
        "I/we the undersigned confirm that I/we have read and agree to abide by the "
        "University regulations on plagiarism and cheating and Faculty coursework policies "
        "and procedures. I/we confirm that this piece of work is my/our own. I/we consent to "
        "appropriate storage of our work for plagiarism checking.")

    doc.add_paragraph("Signature(s): Thong Wai Kit")
    doc.add_paragraph()
    centered("Individual Progressive Report", 14, True, 2)
    centered("Home Credit Loan-Default Prediction", 13, True, 2)
    centered("Thong Wai Kit (P23015668)", 12, False, 2)


def add_toc_page(doc):
    h = doc.add_paragraph()
    run = h.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    h.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2 = p.add_run()
    run2._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3 = p.add_run()
    run3._r.append(fld_sep)

    run4 = p.add_run(
        "Right-click this area and choose Update Field to generate the Table of Contents.")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5 = p.add_run()
    run5._r.append(fld_end)


def add_figure(doc, imgkey, caption, width_in):
    path = IMG[imgkey]
    assert os.path.exists(path), "Image vanished: " + path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(min(width_in, 6.0)))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label, _, rest = caption.partition(":")
    r1 = cap.add_run(label + ":")
    r1.bold = True
    r2 = cap.add_run(rest)
    for r in (r1, r2):
        r.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(12)


def emit(doc, items):
    for item in items:
        kind = item[0]
        if kind == "h1":
            doc.add_paragraph(item[1], style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(item[1], style="Heading 2")
        elif kind == "p":
            p = doc.add_paragraph(item[1])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "li":
            doc.add_paragraph(item[1], style="List Bullet")
        elif kind == "fig":
            imgkey, caption, width = item[1]
            add_figure(doc, imgkey, caption, width)
        else:
            raise ValueError("Unknown item kind: " + kind)


def add_references(doc):
    doc.add_paragraph("References", style="Heading 1")
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(8)
        pf.line_spacing = 1.5


def add_turnitin_page(doc):
    doc.add_paragraph("Turnitin Similarity Percentage Page", style="Heading 1")
    doc.add_paragraph(
        "Insert the Turnitin similarity report screenshot (receipt and percentage page) "
        "here before final submission.")


# ----------------------------------------------------------------------------
# Build
# ----------------------------------------------------------------------------

def body_word_count():
    words = 0
    for chapter in (CH1, CH2, CH3, CH4):
        for item in chapter:
            if item[0] in ("p", "li"):
                words += len(item[1].split())
    return words


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    set_base_styles(doc)

    add_cover_sheet(doc)
    add_page_break(doc)
    add_toc_page(doc)
    add_page_break(doc)

    emit(doc, CH1)
    add_page_break(doc)
    emit(doc, CH2)
    add_page_break(doc)
    emit(doc, CH3)
    add_page_break(doc)
    emit(doc, CH4)
    add_page_break(doc)

    add_references(doc)
    add_page_break(doc)
    emit(doc, APPENDICES)
    add_page_break(doc)
    add_turnitin_page(doc)

    doc.save(OUT_DOCX)
    print("Saved:", OUT_DOCX)
    print("Body word count (Ch1-Ch4 paragraphs and list items):", body_word_count())


if __name__ == "__main__":
    main()
